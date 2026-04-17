"""
gen_informe_comparativo.py
Genera informe comparativo (Word) de carga PM actual vs propuesta.
Output: Asignación Proyectos/3. Output/Informe_Comparativo_Carga_PM_Blend360.docx
"""

import os, copy
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl

# ── Rutas ────────────────────────────────────────────────────────────────────
# Locate "Asignación Proyectos" folder relative to this script OR cwd
_here = os.path.abspath(os.path.dirname(__file__) if '__file__' in dir() else os.getcwd())
# Walk up until we find the folder or use cwd
def _find_asig(start):
    candidate = os.path.join(start, 'Asignación Proyectos')
    if os.path.isdir(candidate):
        return candidate
    parent = os.path.dirname(start)
    if parent == start:
        return None
    return _find_asig(parent)

_asig = _find_asig(_here) or _find_asig(os.getcwd())
if _asig is None:
    # Fallback: hardcode based on known repo root
    _asig = os.path.join(os.path.expanduser('~'), 'Documents', 'PMO-Operations', 'Asignación Proyectos')

BASE = _asig
TEMPLATE   = os.path.join(BASE, '1. Input', 'Plantilla word Blend.docx')
ACTUAL_XL  = os.path.join(BASE, '3. Output', 'Tabla_Carga_PM_Actual_Blend360.xlsx')
PROPUESTA_XL = os.path.join(BASE, '3. Output', 'Propuesta_Reasignacion_PM_Blend360.xlsx')
OUTPUT     = os.path.join(BASE, '3. Output', 'Informe_Comparativo_Carga_PM_Blend360.docx')

# ── Colores Blend ─────────────────────────────────────────────────────────────
BLEND_DARK   = RGBColor(0x00, 0x30, 0x59)   # #003059
BLEND_ACCENT = RGBColor(0x00, 0x70, 0xC0)   # azul medio
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_LIGHT   = RGBColor(0xF2, 0xF2, 0xF2)
GRAY_MED     = RGBColor(0xD9, 0xD9, 0xD9)
RED_SOFT     = RGBColor(0xFF, 0xCC, 0xCC)
YELLOW_SOFT  = RGBColor(0xFF, 0xF2, 0xCC)
GREEN_SOFT   = RGBColor(0xE2, 0xEF, 0xDA)
ORANGE_SOFT  = RGBColor(0xFF, 0xE0, 0xB2)

# ── Datos actuales ────────────────────────────────────────────────────────────
def load_actual():
    wb = openpyxl.load_workbook(ACTUAL_XL)
    ws = wb.active
    rows = []
    header_found = False
    for row in ws.iter_rows(values_only=True):
        if row[0] == 'Project Manager':
            header_found = True
            continue
        if header_found and row[0] and row[0] != 'TOTAL':
            rows.append(row)
        if row[0] == 'TOTAL':
            break
    return rows  # (PM, n_activos, horas_marzo, pct_marzo, estado, horas_abril, pct_abril, reasig, siglas)

def load_propuesta_resumen():
    wb = openpyxl.load_workbook(PROPUESTA_XL)
    ws = wb['Resumen por PM']
    rows = []
    header_found = False
    for row in ws.iter_rows(values_only=True):
        if row[0] == 'Project Manager':
            header_found = True
            continue
        if header_found and row[0] and row[0] != 'TOTAL':
            rows.append(row)
        if row[0] == 'TOTAL':
            break
    return rows  # (PM, proy_act, proy_prop, delta, horas_est, pct_cap, estado, reasig, siglas)

def load_cambios():
    wb = openpyxl.load_workbook(PROPUESTA_XL)
    ws = wb['Cambios de Reasignación']
    rows = []
    header_found = False
    for row in ws.iter_rows(values_only=True):
        if row[0] == 'ID Proyecto':
            header_found = True
            continue
        if header_found and row[0]:
            rows.append(row)  # (id, nombre, cliente, pm_ant, pm_nuevo, horas, motivo)
    return rows

# ── Helpers de formato ────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = str(rgb)  # RGBColor.__str__ returns 6-char hex like '003059'
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, border_color='CCCCCC', size='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def para_fmt(para, font_name='Calibri', size=9, bold=False, color=None,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0):
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    for run in para.runs:
        run.font.name  = font_name
        run.font.size  = Pt(size)
        run.font.bold  = bold
        if color:
            run.font.color.rgb = color

def add_run(para, text, font_name='Calibri', size=9, bold=False, color=None):
    run = para.add_run(text)
    run.font.name  = font_name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = color
    return run

def estado_color(estado):
    if not estado:
        return GRAY_LIGHT
    e = str(estado).lower()
    if 'sobrecargado' in e:
        return RED_SOFT
    if 'alta' in e:
        return ORANGE_SOFT
    if 'normal' in e:
        return GREEN_SOFT
    if 'baja' in e:
        return GRAY_LIGHT
    return GRAY_LIGHT

# ── Tabla con cabecera coloreada ──────────────────────────────────────────────
def make_header_row(table, headers, widths_cm=None):
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, BLEND_DARK)
        set_cell_border(cell, '003059', '4')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(h)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(8)
        run.font.bold  = True
        run.font.color.rgb = WHITE

def fill_data_row(row, values, aligns=None, bg=None, size=8.5):
    for i, val in enumerate(values):
        cell = row.cells[i]
        if bg:
            set_cell_bg(cell, bg)
        else:
            set_cell_bg(cell, GRAY_LIGHT if (row._tr.getparent().index(row._tr) % 2 == 0) else RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_border(cell, 'CCCCCC', '4')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        align = WD_ALIGN_PARAGRAPH.CENTER if (aligns and aligns[i] == 'c') else WD_ALIGN_PARAGRAPH.LEFT
        p.alignment = align
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(str(val) if val is not None else '—')
        run.font.name = 'Calibri'
        run.font.size = Pt(size)

# ── Construcción del documento ────────────────────────────────────────────────
def build_report():
    actual_rows    = load_actual()
    propuesta_rows = load_propuesta_resumen()
    cambios_rows   = load_cambios()

    doc = Document(TEMPLATE)

    # Ajustar márgenes (reducir un poco para que quepan las tablas)
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Limpiar contenido existente del template (conservar header/footer)
    for p in doc.paragraphs[:]:
        p._element.getparent().remove(p._element)
    for t in doc.tables[:]:
        t._element.getparent().remove(t._element)

    body = doc.element.body

    # ── TÍTULO ────────────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    body.append(p_title._element)
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(4)
    r = p_title.add_run('Informe de Carga y Propuesta de Reasignación de Project Managers')
    r.font.name  = 'Calibri'
    r.font.size  = Pt(15)
    r.font.bold  = True
    r.font.color.rgb = BLEND_DARK

    # Subtítulo
    p_sub = doc.add_paragraph()
    body.append(p_sub._element)
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after  = Pt(8)
    r2 = p_sub.add_run('Blend360 Colombia  |  Abril 2026  |  PMO')
    r2.font.name  = 'Calibri'
    r2.font.size  = Pt(9)
    r2.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    # ── INTRODUCCIÓN ─────────────────────────────────────────────────────────
    p_intro = doc.add_paragraph()
    body.append(p_intro._element)
    p_intro.paragraph_format.space_before = Pt(0)
    p_intro.paragraph_format.space_after  = Pt(6)
    add_run(p_intro,
        'Durante marzo de 2026, el análisis de carga de los Project Managers de Blend360 Colombia '
        'evidenció una situación de sobrecarga generalizada: 9 de los 10 PMs activos registraron '
        'horas cercanas o por encima de la capacidad máxima mensual (193,6 h/mes). '
        'El total de horas de gestión registradas en OpenAir para ese período fue de ',
        size=9)
    add_run(p_intro, '1.759,5 h', size=9, bold=True)
    add_run(p_intro,
        ', distribuidas entre cuentas de distintos sectores. '
        'A partir de este diagnóstico, la PMO elaboró una propuesta de reasignación orientada a '
        'equilibrar la carga, alinear portafolios por especialidad y liberar capacidad operativa.',
        size=9)

    # ── SECCIÓN 1: CARGA ACTUAL ───────────────────────────────────────────────
    p_h1 = doc.add_paragraph()
    body.append(p_h1._element)
    p_h1.paragraph_format.space_before = Pt(6)
    p_h1.paragraph_format.space_after  = Pt(3)
    r_h1 = p_h1.add_run('1. Carga actual por Project Manager — Marzo 2026')
    r_h1.font.name  = 'Calibri'
    r_h1.font.size  = Pt(11)
    r_h1.font.bold  = True
    r_h1.font.color.rgb = BLEND_DARK

    # Tabla actual: PM | Proy. Activos | Horas Marzo | % Cap | Estado | Siglas
    headers_act = ['Project Manager', 'Proy.\nActivos', 'Horas\nMarzo', '% Cap.', 'Estado', 'Siglas']
    n_cols_act = len(headers_act)
    tbl_act = doc.add_table(rows=1 + len(actual_rows) + 1, cols=n_cols_act)
    body.append(tbl_act._element)
    tbl_act.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_act.style = 'Table Grid'

    # Anchos columnas (total ~15.5 cm usable)
    col_w_act = [3.8, 1.2, 1.5, 1.3, 2.2, 5.5]
    for c, w in enumerate(col_w_act):
        for row in tbl_act.rows:
            row.cells[c].width = Cm(w)

    make_header_row(tbl_act, headers_act)

    for i, rd in enumerate(actual_rows):
        row = tbl_act.rows[i + 1]
        pm, n_act, h_mar, pct_mar, estado, h_abr, pct_abr, reasig, siglas = rd
        pct_str = f'{pct_mar*100:.1f}%' if pct_mar else '0%'
        estado_clean = str(estado).replace(' ⚠', '') if estado else '—'
        bg = estado_color(estado)
        vals  = [pm, str(n_act), f'{h_mar:.0f}h', pct_str, estado_clean, siglas or '—']
        aligns = ['l', 'c', 'c', 'c', 'c', 'l']
        fill_data_row(row, vals, aligns, bg=None, size=8.5)
        # Estado con color
        est_cell = row.cells[4]
        set_cell_bg(est_cell, bg)

    # Fila TOTAL
    row_tot = tbl_act.rows[len(actual_rows) + 1]
    set_cell_bg(row_tot.cells[0], BLEND_DARK)
    set_cell_border(row_tot.cells[0], '003059', '4')
    p0 = row_tot.cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p0.paragraph_format.space_before = Pt(1)
    p0.paragraph_format.space_after  = Pt(1)
    r0 = p0.add_run('TOTAL')
    r0.font.name = 'Calibri'; r0.font.size = Pt(8.5); r0.font.bold = True; r0.font.color.rgb = WHITE

    set_cell_bg(row_tot.cells[2], BLEND_DARK)
    set_cell_border(row_tot.cells[2], '003059', '4')
    p2 = row_tot.cells[2].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after  = Pt(1)
    r2t = p2.add_run('1.759,5h')
    r2t.font.name = 'Calibri'; r2t.font.size = Pt(8.5); r2t.font.bold = True; r2t.font.color.rgb = WHITE

    for ci in [1, 3, 4, 5]:
        set_cell_bg(row_tot.cells[ci], BLEND_DARK)
        set_cell_border(row_tot.cells[ci], '003059', '4')
        row_tot.cells[ci].paragraphs[0].paragraph_format.space_before = Pt(1)
        row_tot.cells[ci].paragraphs[0].paragraph_format.space_after  = Pt(1)

    # Nota David Cortes
    p_nota = doc.add_paragraph()
    body.append(p_nota._element)
    p_nota.paragraph_format.space_before = Pt(3)
    p_nota.paragraph_format.space_after  = Pt(3)
    add_run(p_nota, '* ', size=8, bold=True, color=BLEND_ACCENT)
    add_run(p_nota,
        'David Cortes: incluye el proyecto FDN (Financiera de Desarrollo Nacional, P2030) que '
        'a la fecha de análisis aún no había iniciado ejecución. Descontando ese proyecto, '
        'su carga real de marzo se aproxima a 192 h (99 % de capacidad).',
        size=8, color=RGBColor(0x40, 0x40, 0x40))

    # ── SECCIÓN 2: PROPUESTA ──────────────────────────────────────────────────
    p_h2 = doc.add_paragraph()
    body.append(p_h2._element)
    p_h2.paragraph_format.space_before = Pt(8)
    p_h2.paragraph_format.space_after  = Pt(3)
    r_h2 = p_h2.add_run('2. Propuesta de reasignación — Carga estimada post-ajuste')
    r_h2.font.name  = 'Calibri'
    r_h2.font.size  = Pt(11)
    r_h2.font.bold  = True
    r_h2.font.color.rgb = BLEND_DARK

    p_intro2 = doc.add_paragraph()
    body.append(p_intro2._element)
    p_intro2.paragraph_format.space_before = Pt(0)
    p_intro2.paragraph_format.space_after  = Pt(4)
    add_run(p_intro2,
        'La propuesta redistribuye 25 sub-proyectos entre 8 movimientos de PM, reduciendo '
        'el total de horas estimadas de gestión a ',
        size=9)
    add_run(p_intro2, '1.402,7 h/mes', size=9, bold=True)
    add_run(p_intro2,
        '. Se eliminan los casos de sobrecarga crítica (excepción: Oscar Barragan cuyo portafolio '
        'AGR/BTG/UCA/UNA/PRO/CAT no tiene alternativa de reasignación inmediata), y se logra '
        'alineación de portafolios por especialidad.',
        size=9)

    # Tabla propuesta: PM | Proy. Prop | Δ | Horas Est | % Cap | Estado | Siglas
    headers_prop = ['Project Manager', 'Proy.\nProp.', 'Δ', 'Horas\nEst./mes', '% Cap.', 'Estado', 'Siglas']
    n_cols_prop = len(headers_prop)
    tbl_prop = doc.add_table(rows=1 + len(propuesta_rows) + 1, cols=n_cols_prop)
    body.append(tbl_prop._element)
    tbl_prop.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_prop.style = 'Table Grid'

    col_w_prop = [3.5, 1.2, 0.9, 1.6, 1.3, 2.2, 4.8]
    for c, w in enumerate(col_w_prop):
        for row in tbl_prop.rows:
            row.cells[c].width = Cm(w)

    make_header_row(tbl_prop, headers_prop)

    for i, rd in enumerate(propuesta_rows):
        row = tbl_prop.rows[i + 1]
        pm, proy_act, proy_prop, delta, horas_est, pct_cap, estado, reasig, siglas = rd
        delta_str = ('+' if delta > 0 else '') + str(int(delta)) if delta else '0'
        pct_str = f'{pct_cap*100:.1f}%' if pct_cap else '0%'
        estado_clean = str(estado).replace(' ⚠', '') if estado else '—'
        bg = estado_color(estado)
        vals = [pm, str(int(proy_prop)) if proy_prop else '—', delta_str,
                f'{horas_est:.0f}h', pct_str, estado_clean, siglas or '—']
        aligns = ['l', 'c', 'c', 'c', 'c', 'c', 'l']
        fill_data_row(row, vals, aligns, bg=None, size=8.5)
        set_cell_bg(row.cells[5], bg)

    # Fila TOTAL propuesta
    row_tot2 = tbl_prop.rows[len(propuesta_rows) + 1]
    for ci in range(n_cols_prop):
        set_cell_bg(row_tot2.cells[ci], BLEND_DARK)
        set_cell_border(row_tot2.cells[ci], '003059', '4')
        row_tot2.cells[ci].paragraphs[0].paragraph_format.space_before = Pt(1)
        row_tot2.cells[ci].paragraphs[0].paragraph_format.space_after  = Pt(1)
    p0t = row_tot2.cells[0].paragraphs[0]
    p0t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rt0 = p0t.add_run('TOTAL')
    rt0.font.name = 'Calibri'; rt0.font.size = Pt(8.5); rt0.font.bold = True; rt0.font.color.rgb = WHITE

    p3t = row_tot2.cells[3].paragraphs[0]
    p3t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt3 = p3t.add_run('1.402,7h')
    rt3.font.name = 'Calibri'; rt3.font.size = Pt(8.5); rt3.font.bold = True; rt3.font.color.rgb = WHITE

    # ── SECCIÓN 3: CAMBIOS ────────────────────────────────────────────────────
    p_h3 = doc.add_paragraph()
    body.append(p_h3._element)
    p_h3.paragraph_format.space_before = Pt(8)
    p_h3.paragraph_format.space_after  = Pt(3)
    r_h3 = p_h3.add_run('3. Detalle de cambios en la asignación')
    r_h3.font.name  = 'Calibri'
    r_h3.font.size  = Pt(11)
    r_h3.font.bold  = True
    r_h3.font.color.rgb = BLEND_DARK

    # Agrupar cambios por (sigla, cliente, pm_ant, pm_nuevo, motivo)
    from collections import defaultdict
    grupos = defaultdict(lambda: {'ids': [], 'horas': 0, 'cliente': '', 'motivo': ''})
    for row in cambios_rows:
        id_p, nombre, cliente, pm_ant, pm_nuevo, horas, motivo = row
        sigla = nombre.split(' - ')[0].strip() if nombre else ''
        key = (sigla, pm_ant, pm_nuevo)
        grupos[key]['ids'].append(id_p)
        grupos[key]['horas'] = horas  # mismas para todos los sub-proyectos de una sigla
        grupos[key]['cliente'] = (cliente or '').strip()
        grupos[key]['motivo'] = motivo or ''

    headers_ch = ['Sigla / Cliente', 'Sub-proy.', 'PM Anterior', 'PM Nuevo', 'h/mes est.', 'Criterio']
    tbl_ch = doc.add_table(rows=1 + len(grupos), cols=len(headers_ch))
    body.append(tbl_ch._element)
    tbl_ch.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_ch.style = 'Table Grid'

    col_w_ch = [3.8, 1.2, 2.2, 2.2, 1.4, 4.7]
    for c, w in enumerate(col_w_ch):
        for row in tbl_ch.rows:
            row.cells[c].width = Cm(w)

    make_header_row(tbl_ch, headers_ch)

    for i, ((sigla, pm_ant, pm_nuevo), info) in enumerate(sorted(grupos.items(), key=lambda x: x[0][2])):
        row = tbl_ch.rows[i + 1]
        n_sub = len(info['ids'])
        total_h = info['horas'] * n_sub
        cliente_short = info['cliente'][:30] + ('…' if len(info['cliente']) > 30 else '')
        sigla_cliente = f'{sigla}\n{cliente_short}'
        motivo_short = info['motivo'][:55] + ('…' if len(info['motivo']) > 55 else '')
        vals = [sigla_cliente, str(n_sub), pm_ant, pm_nuevo, f'{total_h:.0f}h', motivo_short]
        aligns = ['l', 'c', 'l', 'l', 'c', 'l']
        fill_data_row(row, vals, aligns, bg=None, size=8)

    # ── CONCLUSIÓN ────────────────────────────────────────────────────────────
    p_conc = doc.add_paragraph()
    body.append(p_conc._element)
    p_conc.paragraph_format.space_before = Pt(8)
    p_conc.paragraph_format.space_after  = Pt(2)
    r_conc = p_conc.add_run('Conclusiones y próximos pasos')
    r_conc.font.name  = 'Calibri'
    r_conc.font.size  = Pt(11)
    r_conc.font.bold  = True
    r_conc.font.color.rgb = BLEND_DARK

    bullets = [
        ('Reducción de carga total: ', 'de 1.759,5 h a 1.402,7 h/mes estimadas (–20,3 %), con 7 de 9 PMs por debajo del 100 % de capacidad.'),
        ('Alineación por portafolio: ', 'educación (SED/MEN) consolidado en Diana Castro; datos/observatorios (SPD) en Diana Rojas; cloud managed services (HPL, JPM) en David Cortes y Juan Bernal.'),
        ('Riesgos residuales: ', 'Oscar Barragan permanece con carga alta (111 %). Se recomienda evaluar soporte de gestión o contratación de un PM adicional para el portafolio AGR/BTG. David Cortes (138 % est.) se normaliza una vez inicie FDN; se monitorea en mayo.'),
        ('Acción inmediata: ', 'actualizar maestro.proyectos.xlsx con los nuevos PMs asignados y comunicar los cambios a los equipos de proyecto antes del 25 de abril de 2026.'),
    ]
    for bold_txt, normal_txt in bullets:
        p_b = doc.add_paragraph()
        body.append(p_b._element)
        p_b.paragraph_format.space_before = Pt(1)
        p_b.paragraph_format.space_after  = Pt(1)
        p_b.paragraph_format.left_indent  = Cm(0.5)
        add_run(p_b, '• ', size=9, bold=True, color=BLEND_ACCENT)
        add_run(p_b, bold_txt, size=9, bold=True)
        add_run(p_b, normal_txt, size=9)

    doc.save(OUTPUT)
    print(f'OK Informe guardado en: {OUTPUT}')

if __name__ == '__main__':
    build_report()
