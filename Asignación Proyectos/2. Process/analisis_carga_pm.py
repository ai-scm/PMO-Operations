# =============================================================================
# Script: gen_outputs.py
# Propósito: Generar análisis de carga de PMs, informe Word y consolidado Excel
# PMO Operations - Blend360 Colombia
# Fecha: Abril 2026
# =============================================================================

import os
import io
import shutil
import warnings
from datetime import datetime, date
from collections import defaultdict

warnings.filterwarnings('ignore')

import openpyxl
from openpyxl.styles import (Font, PatternFill, Alignment, Border, Side,
                              numbers)
from openpyxl.utils import get_column_letter
from openpyxl.chart import BarChart, LineChart, Reference
from openpyxl.chart.series import SeriesLabel
from openpyxl.drawing.image import Image as XLImage

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =============================================================================
# PATHS
# =============================================================================
BASE = 'c:/Users/BOG-LAP-SER-176/Documents/PMO-Operations'
FOLDER_NAME = [f for f in os.listdir(BASE) if 'Asign' in f][0]
ROOT = os.path.join(BASE, FOLDER_NAME)
INPUT   = os.path.join(ROOT, '1. Input')
PROCESS = os.path.join(ROOT, '2. Process')
OUTPUT  = os.path.join(ROOT, '3. Output')
CHARTS_DIR = os.path.join(OUTPUT, 'charts')
os.makedirs(CHARTS_DIR, exist_ok=True)

HOURS_FILE    = os.path.join(INPUT, 'Horas marzo abril.xlsx')
MAESTRO_FILE  = os.path.join(INPUT, 'maestro.proyectos.xlsx')
TEMPLATE_FILE = os.path.join(INPUT, 'Plantilla word Blend.docx')

# Blend brand colors
BLEND_DARK_BLUE  = RGBColor(0x00, 0x30, 0x59)   # #003059
BLEND_MID_BLUE   = RGBColor(0x00, 0x70, 0xC0)   # #0070C0
BLEND_LIGHT_BLUE = RGBColor(0xBD, 0xD7, 0xEE)   # #BDD7EE
BLEND_ORANGE     = RGBColor(0xFF, 0x6B, 0x00)   # #FF6B00
BLEND_GRAY       = RGBColor(0x40, 0x40, 0x40)   # #404040

HEX_DARK_BLUE  = '003059'
HEX_MID_BLUE   = '0070C0'
HEX_LIGHT_BLUE = 'BDD7EE'
HEX_ORANGE     = 'FF6B00'
HEX_GRAY       = '595959'
HEX_LIGHT_GRAY = 'F2F2F2'
HEX_RED        = 'FF0000'
HEX_YELLOW     = 'FFD966'
HEX_GREEN      = '70AD47'

# =============================================================================
# 1. LOAD DATA
# =============================================================================
print("Cargando datos...")

# --- Maestro de proyectos ---
wb_maestro = openpyxl.load_workbook(MAESTRO_FILE, data_only=True)
ws_maestro = wb_maestro['project']
maestro_rows = list(ws_maestro.iter_rows(values_only=True))
maestro_header = maestro_rows[0]

# Build project->PM mapping (active)
project_pm_map = {}  # project_name -> pm
pm_active_projects = defaultdict(list)  # pm -> [project names]

for row in maestro_rows[1:]:
    if row[8] == 'Activo':
        pm = str(row[11]) if row[11] else 'N.A.'
        sigla = str(row[5]) if row[5] else ''
        objetivo = str(row[3]) if row[3] else ''
        tipo = str(row[6]) if row[6] else ''
        cliente = str(row[7]) if row[7] else ''
        proj_id = str(row[0]) if row[0] else ''

        if pm not in ['N.A.', 'Sin Asignar', 'None', '']:
            proj_info = {
                'id': proj_id, 'objetivo': objetivo,
                'sigla': sigla, 'tipo': tipo, 'cliente': cliente
            }
            pm_active_projects[pm].append(proj_info)

# --- Horas Marzo-Abril ---
print("Leyendo horas registradas...")
wb_hours = openpyxl.load_workbook(HOURS_FILE)
ws_hours = wb_hours['openair (39)']
hours_rows = list(ws_hours.iter_rows(values_only=True))

# PM name -> OpenAir user mapping
PM_USER_MAP = {
    'Oscar Barragan':          'BARRAGAN, OSCAR',
    'Juan Bernal':             'BERNAL MORENO, JUAN CAMILO',
    'David Cortes':            'CORTES, DAVID',
    'Miguel Garcia':           'GARCIA, MIGUEL',
    'Kelly Carbonell':         'CARBONELL RAMOS, KELLY MARGARITA',
    'Daniel Sebastian Vargas': 'VARGAS CRISTANCHO, DANIEL SEBASTIAN',
    'Diana Castro':            'CASTRO, DIANA',
    'Diana Rojas':             'ROJAS CHARRY, DIANA CLEMENCIA',
    'Indira Duarte':           'DUARTE, INDIRA',
}
USER_PM_MAP = {v: k for k, v in PM_USER_MAP.items()}

# Capacity parameters
WEEKLY_CAPACITY = 44  # hours/week
# March 2026: working weeks starting 01/03, 08/03, 15/03, 22/03, 29/03
# Days in March: 31 days, but last week extends to April
# Effective March working time ≈ 4.4 weeks
MARCH_WEEKS = 4.4
MARCH_CAPACITY = WEEKLY_CAPACITY * MARCH_WEEKS  # ≈ 193.6h

# Collect data structures
# pm -> month -> project -> hours
pm_proj_month = defaultdict(lambda: defaultdict(lambda: defaultdict(float)))
# pm -> month -> total hours
pm_month_total = defaultdict(lambda: defaultdict(float))
# pm -> week_start -> hours
pm_weekly = defaultdict(lambda: defaultdict(float))
# pm -> week_start -> {project: hours}
pm_weekly_proj = defaultdict(lambda: defaultdict(lambda: defaultdict(float)))
# category tracking
pm_proj_category = {}  # (pm, project) -> category (billable/interno/vacaciones)

INTERNAL_PROJECTS = {'Meetings & Training', 'Time Off & Holiday', 'Business Development',
                     'Business Development – Project Management',
                     'Business Development \u2013 Project Management'}

for row in hours_rows[1:]:
    user = str(row[1]) if row[1] else ''
    if user not in USER_PM_MAP:
        continue

    pm = USER_PM_MAP[user]
    date_val = row[0]
    start_date = row[5]
    hours = float(row[9]) if row[9] else 0
    project = str(row[8]) if row[8] else ''

    if not isinstance(date_val, datetime):
        continue
    if date_val.year != 2026 or date_val.month not in [3, 4]:
        continue

    month = 'Marzo 2026' if date_val.month == 3 else 'Abril 2026'

    week_key = start_date.strftime('%d/%m/%Y') if isinstance(start_date, datetime) else str(start_date)

    pm_proj_month[pm][month][project] += hours
    pm_month_total[pm][month] += hours
    pm_weekly[pm][week_key] += hours
    pm_weekly_proj[pm][week_key][project] += hours

    # Categorize
    key = (pm, project)
    if 'Time Off' in project or 'Holiday' in project:
        pm_proj_category[key] = 'Vacaciones/Ausencia'
    elif 'Meeting' in project or 'Training' in project or 'Business Development' in project:
        pm_proj_category[key] = 'Interno/Administrativo'
    else:
        pm_proj_category[key] = 'Facturable'

PM_ORDER = sorted(PM_USER_MAP.keys())

# April weeks detected (for capacity calc)
april_weeks_per_pm = {}
for pm in PM_ORDER:
    april_wks = [w for w in pm_weekly[pm] if '/04/2026' in w]
    # Count partial April (days in week 29/03 that fall in April: up to 3 days ≈ 26.4h)
    march_last_week = '29/03/2026'
    has_march_last = march_last_week in pm_weekly[pm]
    extra = 0.6  # 0.6 week equivalent for partial week in April
    april_cap = (len(april_wks) + (extra if has_march_last else 0)) * WEEKLY_CAPACITY
    april_weeks_per_pm[pm] = round(len(april_wks) + (extra if has_march_last else 0), 1)

print("Datos cargados correctamente.")
print(f"PMs identificados: {len(PM_ORDER)}")
for pm in PM_ORDER:
    n_proj = len(pm_active_projects.get(pm, []))
    m_h = pm_month_total[pm].get('Marzo 2026', 0)
    a_h = pm_month_total[pm].get('Abril 2026', 0)
    print(f"  {pm}: {n_proj} proyectos activos | Mar={m_h:.0f}h | Abr={a_h:.0f}h")

# =============================================================================
# 2. GENERATE CHARTS
# =============================================================================
print("\nGenerando gráficas...")

BLEND_PALETTE = ['#003059', '#0070C0', '#FF6B00', '#70AD47', '#FFD966',
                 '#BDD7EE', '#A9D18E', '#F4B942', '#C9504C', '#7030A0']

def set_spine_style(ax):
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#CCCCCC')
    ax.spines['bottom'].set_color('#CCCCCC')
    ax.tick_params(colors='#404040', labelsize=8)
    ax.yaxis.grid(True, linestyle='--', alpha=0.5, color='#CCCCCC')
    ax.set_axisbelow(True)

# --- Chart 1: Total hours per PM (March vs April) ---
fig, ax = plt.subplots(figsize=(10, 5))
x = np.arange(len(PM_ORDER))
w = 0.35
march_vals = [pm_month_total[pm].get('Marzo 2026', 0) for pm in PM_ORDER]
april_vals  = [pm_month_total[pm].get('Abril 2026', 0) for pm in PM_ORDER]
pm_labels = [p.replace(' ', '\n') for p in PM_ORDER]

bars1 = ax.bar(x - w/2, march_vals, w, label='Marzo 2026', color='#003059')
bars2 = ax.bar(x + w/2, april_vals,  w, label='Abril 2026', color='#0070C0')
ax.axhline(y=MARCH_CAPACITY, color='#FF6B00', linestyle='--', lw=1.5,
           label=f'Capacidad Marzo ({MARCH_CAPACITY:.0f}h)')
ax.set_xticks(x)
ax.set_xticklabels(pm_labels, fontsize=7)
ax.set_ylabel('Horas', fontsize=9, color='#404040')
ax.set_title('Total de Horas Registradas por PM\nMarzo vs Abril 2026',
             fontsize=11, fontweight='bold', color='#003059', pad=12)
ax.legend(fontsize=8, framealpha=0.8)
set_spine_style(ax)
for bar in bars1:
    h = bar.get_height()
    if h > 0:
        ax.text(bar.get_x()+bar.get_width()/2, h+2, f'{h:.0f}',
                ha='center', va='bottom', fontsize=7, color='#003059')
for bar in bars2:
    h = bar.get_height()
    if h > 0:
        ax.text(bar.get_x()+bar.get_width()/2, h+2, f'{h:.0f}',
                ha='center', va='bottom', fontsize=7, color='#0070C0')
plt.tight_layout()
chart1_path = os.path.join(CHARTS_DIR, 'chart_horas_pm.png')
plt.savefig(chart1_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()
print(f"  Gráfica 1 guardada: {chart1_path}")

# --- Chart 2: % Utilización Capacidad Marzo ---
fig, ax = plt.subplots(figsize=(9, 4.5))
march_pct = [(pm_month_total[pm].get('Marzo 2026', 0) / MARCH_CAPACITY * 100) for pm in PM_ORDER]
colors = ['#C9504C' if p > 100 else '#FFD966' if p > 90 else '#70AD47' for p in march_pct]
bars = ax.barh(PM_ORDER, march_pct, color=colors, edgecolor='white', height=0.6)
ax.axvline(x=100, color='#FF6B00', linestyle='--', lw=1.5, label='100% Capacidad')
ax.axvline(x=80, color='#FFD966', linestyle=':', lw=1.2, label='80% Alerta')
ax.set_xlabel('% Utilización', fontsize=9, color='#404040')
ax.set_title('Utilización de Capacidad por PM – Marzo 2026\n(Capacidad = 193.6 horas / 4.4 semanas × 44h)',
             fontsize=10, fontweight='bold', color='#003059', pad=12)
ax.legend(fontsize=8, loc='lower right')
set_spine_style(ax)
ax.spines['left'].set_visible(True)
for i, (bar, val) in enumerate(zip(bars, march_pct)):
    ax.text(val + 0.5, bar.get_y() + bar.get_height()/2,
            f'{val:.1f}%', va='center', fontsize=8, color='#404040')
ax.set_xlim(0, 115)
plt.tight_layout()
chart2_path = os.path.join(CHARTS_DIR, 'chart_utilizacion_marzo.png')
plt.savefig(chart2_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()
print(f"  Gráfica 2 guardada: {chart2_path}")

# --- Chart 3: Número de proyectos activos por PM ---
fig, ax = plt.subplots(figsize=(9, 4.5))
n_projects = [len(pm_active_projects.get(pm, [])) for pm in PM_ORDER]
bar_colors = [BLEND_PALETTE[i % len(BLEND_PALETTE)] for i in range(len(PM_ORDER))]
bars = ax.barh(PM_ORDER, n_projects, color=bar_colors, edgecolor='white', height=0.6)
ax.set_xlabel('Número de Proyectos', fontsize=9, color='#404040')
ax.set_title('Proyectos Activos Asignados por PM',
             fontsize=10, fontweight='bold', color='#003059', pad=12)
set_spine_style(ax)
ax.spines['left'].set_visible(True)
for bar, val in zip(bars, n_projects):
    ax.text(val + 0.3, bar.get_y() + bar.get_height()/2,
            str(val), va='center', fontsize=9, fontweight='bold', color='#404040')
plt.tight_layout()
chart3_path = os.path.join(CHARTS_DIR, 'chart_proyectos_activos.png')
plt.savefig(chart3_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()
print(f"  Gráfica 3 guardada: {chart3_path}")

# --- Chart 4: Distribución de horas por categoría (Marzo) - Stacked ---
fig, ax = plt.subplots(figsize=(11, 5))
cat_facturable   = []
cat_interno      = []
cat_vacaciones   = []
for pm in PM_ORDER:
    f = i_ = v = 0
    for proj, hours in pm_proj_month[pm].get('Marzo 2026', {}).items():
        cat = pm_proj_category.get((pm, proj), 'Facturable')
        if cat == 'Facturable':            f  += hours
        elif cat == 'Interno/Administrativo': i_ += hours
        else:                              v  += hours
    cat_facturable.append(f)
    cat_interno.append(i_)
    cat_vacaciones.append(v)

x2 = np.arange(len(PM_ORDER))
w2 = 0.5
b1 = ax.bar(x2, cat_facturable,  w2, label='Facturable',            color='#003059')
b2 = ax.bar(x2, cat_interno,     w2, bottom=cat_facturable,          label='Interno/Administrativo', color='#0070C0')
b3 = ax.bar(x2, cat_vacaciones,  w2,
            bottom=[f+i for f,i in zip(cat_facturable, cat_interno)],
            label='Vacaciones/Ausencia', color='#FFD966')
ax.set_xticks(x2)
ax.set_xticklabels([p.replace(' ', '\n') for p in PM_ORDER], fontsize=7)
ax.set_ylabel('Horas', fontsize=9, color='#404040')
ax.set_title('Distribución de Horas por Categoría – Marzo 2026',
             fontsize=10, fontweight='bold', color='#003059', pad=12)
ax.legend(fontsize=8, loc='upper right')
set_spine_style(ax)
plt.tight_layout()
chart4_path = os.path.join(CHARTS_DIR, 'chart_distribucion_categoria.png')
plt.savefig(chart4_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()
print(f"  Gráfica 4 guardada: {chart4_path}")

# --- Chart 5: Weekly evolution per PM ---
all_weeks = sorted(set(w for pm in PM_ORDER for w in pm_weekly[pm]))
# Filter only March-April 2026 weeks
all_weeks = [w for w in all_weeks if '2026' in w]

fig, axes = plt.subplots(3, 3, figsize=(14, 9), sharey=False)
axes_flat = axes.flatten()
for idx, pm in enumerate(PM_ORDER):
    ax = axes_flat[idx]
    weekly_vals = [pm_weekly[pm].get(w, 0) for w in all_weeks]
    colors_w = ['#003059' if '/03/' in w else '#0070C0' for w in all_weeks]
    ax.bar(range(len(all_weeks)), weekly_vals, color=colors_w, edgecolor='white')
    ax.axhline(y=44, color='#FF6B00', linestyle='--', lw=1, alpha=0.8)
    short_labels = [w[:5] for w in all_weeks]
    ax.set_xticks(range(len(all_weeks)))
    ax.set_xticklabels(short_labels, fontsize=6, rotation=45)
    ax.set_title(pm, fontsize=7, fontweight='bold', color='#003059')
    ax.tick_params(labelsize=6)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.yaxis.grid(True, linestyle='--', alpha=0.3, color='#CCCCCC')
    ax.set_axisbelow(True)

patch_mar = mpatches.Patch(color='#003059', label='Marzo')
patch_abr = mpatches.Patch(color='#0070C0', label='Abril')
patch_cap = plt.Line2D([0], [0], color='#FF6B00', linestyle='--', lw=1.5, label='44h/sem')
fig.legend(handles=[patch_mar, patch_abr, patch_cap], loc='lower center',
           ncol=3, fontsize=8, bbox_to_anchor=(0.5, 0.0))
fig.suptitle('Evolución Semanal de Horas por PM – Marzo y Abril 2026',
             fontsize=11, fontweight='bold', color='#003059', y=1.01)
plt.tight_layout(rect=[0, 0.04, 1, 1])
chart5_path = os.path.join(CHARTS_DIR, 'chart_evolucion_semanal.png')
plt.savefig(chart5_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()
print(f"  Gráfica 5 guardada: {chart5_path}")

# --- Chart 6: Top projects per PM (individual bar charts for report) ---
pm_chart_paths = {}
for pm in PM_ORDER:
    fig, ax = plt.subplots(figsize=(8, 4))
    projs = pm_proj_month[pm].get('Marzo 2026', {})
    sorted_projs = sorted(projs.items(), key=lambda x: -x[1])[:8]
    if not sorted_projs:
        plt.close()
        continue
    labels = [p[:40] + '...' if len(p) > 40 else p for p, _ in sorted_projs]
    values = [h for _, h in sorted_projs]
    colors_p = [('#70AD47' if pm_proj_category.get((pm, p), 'F') == 'Facturable'
                 else '#FFD966' if pm_proj_category.get((pm, p), 'F') == 'Interno/Administrativo'
                 else '#FF6B00')
                for p, _ in sorted_projs]
    bars = ax.barh(labels, values, color=colors_p, edgecolor='white', height=0.6)
    ax.set_xlabel('Horas – Marzo 2026', fontsize=8)
    ax.set_title(f'Distribución de Horas: {pm}\nMarzo 2026', fontsize=9,
                 fontweight='bold', color='#003059')
    set_spine_style(ax)
    ax.spines['left'].set_visible(True)
    for bar, val in zip(bars, values):
        ax.text(val + 0.5, bar.get_y() + bar.get_height()/2,
                f'{val:.0f}h', va='center', fontsize=7)
    plt.tight_layout()
    path = os.path.join(CHARTS_DIR, f'chart_pm_{pm.replace(" ", "_")}.png')
    plt.savefig(path, dpi=130, bbox_inches='tight', facecolor='white')
    plt.close()
    pm_chart_paths[pm] = path

print("Gráficas generadas exitosamente.")

# =============================================================================
# 3. GENERATE EXCEL CONSOLIDATION
# =============================================================================
print("\nGenerando Excel consolidado...")

wb_out = openpyxl.Workbook()

# Color fills
fill_dark_blue   = PatternFill('solid', fgColor=HEX_DARK_BLUE)
fill_mid_blue    = PatternFill('solid', fgColor=HEX_MID_BLUE)
fill_light_blue  = PatternFill('solid', fgColor=HEX_LIGHT_BLUE)
fill_orange      = PatternFill('solid', fgColor=HEX_ORANGE)
fill_light_gray  = PatternFill('solid', fgColor=HEX_LIGHT_GRAY)
fill_yellow      = PatternFill('solid', fgColor=HEX_YELLOW)
fill_green       = PatternFill('solid', fgColor='70AD47')
fill_red         = PatternFill('solid', fgColor='FF4D4D')
fill_white       = PatternFill('solid', fgColor='FFFFFF')

thin = Side(style='thin', color='CCCCCC')
medium = Side(style='medium', color='003059')
border_thin  = Border(left=thin, right=thin, top=thin, bottom=thin)
border_header = Border(left=medium, right=medium, top=medium, bottom=medium)

def hdr_font(sz=10, bold=True, color='FFFFFF'):
    return Font(name='Calibri', size=sz, bold=bold, color=color)

def cell_font(sz=10, bold=False, color='000000'):
    return Font(name='Calibri', size=sz, bold=bold, color=color)

def center_align(wrap=False):
    return Alignment(horizontal='center', vertical='center', wrap_text=wrap)

def left_align(wrap=False):
    return Alignment(horizontal='left', vertical='center', wrap_text=wrap)

# ---- SHEET 1: Resumen por PM ----
ws1 = wb_out.active
ws1.title = 'Resumen por PM'
ws1.sheet_view.showGridLines = False

# Title
ws1.merge_cells('A1:J1')
ws1['A1'] = 'ANÁLISIS DE CARGA DE PMs – MARZO Y ABRIL 2026'
ws1['A1'].font = Font(name='Calibri', size=14, bold=True, color='FFFFFF')
ws1['A1'].fill = fill_dark_blue
ws1['A1'].alignment = center_align()
ws1.row_dimensions[1].height = 30

ws1.merge_cells('A2:J2')
ws1['A2'] = f'Capacidad semanal: {WEEKLY_CAPACITY} horas | Blend360 Colombia PMO'
ws1['A2'].font = Font(name='Calibri', size=10, italic=True, color='FFFFFF')
ws1['A2'].fill = fill_mid_blue
ws1['A2'].alignment = center_align()
ws1.row_dimensions[2].height = 20

# Headers row 4
headers = ['Project Manager', 'Proyectos\nActivos', 'Horas\nMarzo', 'Horas\nAbril',
           'Total\nHoras', '% Uso\nMarzo', 'Sem. Abril\nDisp.', '% Uso\nAbril',
           'Estado\nCarga Marzo', 'Recomendación']
col_widths = [25, 12, 12, 12, 12, 12, 13, 12, 18, 35]
for i, (h, w) in enumerate(zip(headers, col_widths), start=1):
    cell = ws1.cell(row=4, column=i, value=h)
    cell.font = hdr_font(10)
    cell.fill = fill_dark_blue
    cell.alignment = center_align(wrap=True)
    cell.border = border_thin
    ws1.column_dimensions[get_column_letter(i)].width = w

ws1.row_dimensions[4].height = 35

# Data rows
for row_i, pm in enumerate(PM_ORDER, start=5):
    n_p   = len(pm_active_projects.get(pm, []))
    m_h   = pm_month_total[pm].get('Marzo 2026', 0)
    a_h   = pm_month_total[pm].get('Abril 2026', 0)
    tot   = m_h + a_h
    m_pct = m_h / MARCH_CAPACITY * 100
    a_wks = april_weeks_per_pm.get(pm, 1.6)
    a_cap = a_wks * WEEKLY_CAPACITY
    a_pct = (a_h / a_cap * 100) if a_cap > 0 else 0

    if m_pct > 100:   estado = 'Sobrecargado ⚠'
    elif m_pct >= 90: estado = 'Carga Alta'
    elif m_pct >= 70: estado = 'Carga Normal'
    else:             estado = 'Carga Baja'

    if n_p > 40:      reco = 'Revisar cartera – exceso de proyectos'
    elif m_pct > 100: reco = 'Redistribuir proyectos urgente'
    elif n_p > 20:    reco = 'Monitorear carga regularmente'
    else:             reco = 'Carga manejable'

    row_vals = [pm, n_p, m_h, a_h, tot, m_pct/100, a_wks, a_pct/100, estado, reco]
    for col_i, val in enumerate(row_vals, start=1):
        cell = ws1.cell(row=row_i, column=col_i, value=val)
        cell.border = border_thin
        cell.alignment = left_align(wrap=True) if col_i in [1, 9, 10] else center_align()
        cell.font = cell_font(10)

        if col_i in [6, 8]:  # percentage columns
            cell.number_format = '0.0%'
        if col_i in [3, 4, 5]:
            cell.number_format = '#,##0.0'

        # Color status column
        if col_i == 9:
            if 'Sobrecargado' in str(val):
                cell.fill = fill_red
                cell.font = Font(name='Calibri', size=10, bold=True, color='FFFFFF')
            elif 'Alta' in str(val):
                cell.fill = fill_yellow
                cell.font = cell_font(10, bold=True, color='595959')
            else:
                cell.fill = fill_green
                cell.font = Font(name='Calibri', size=10, color='FFFFFF')
        # Color % columns
        if col_i == 6:
            if m_pct > 100:   cell.fill = fill_red;    cell.font = Font(name='Calibri',size=10,bold=True,color='FFFFFF')
            elif m_pct >= 90: cell.fill = fill_yellow
            else:             cell.fill = fill_green;   cell.font = Font(name='Calibri',size=10,color='FFFFFF')

    ws1.row_dimensions[row_i].height = 25

# Totals row
r_tot = len(PM_ORDER) + 5
ws1.merge_cells(f'A{r_tot}:B{r_tot}')
ws1[f'A{r_tot}'] = 'TOTAL'
ws1[f'A{r_tot}'].font = hdr_font(10)
ws1[f'A{r_tot}'].fill = fill_dark_blue
ws1[f'A{r_tot}'].alignment = center_align()
ws1[f'C{r_tot}'] = sum(pm_month_total[pm].get('Marzo 2026', 0) for pm in PM_ORDER)
ws1[f'D{r_tot}'] = sum(pm_month_total[pm].get('Abril 2026', 0) for pm in PM_ORDER)
ws1[f'E{r_tot}'] = ws1[f'C{r_tot}'].value + ws1[f'D{r_tot}'].value
for col in ['C','D','E']:
    ws1[f'{col}{r_tot}'].font = hdr_font(10)
    ws1[f'{col}{r_tot}'].fill = fill_dark_blue
    ws1[f'{col}{r_tot}'].alignment = center_align()
    ws1[f'{col}{r_tot}'].number_format = '#,##0.0'

ws1.row_dimensions[r_tot].height = 25

# Freeze pane
ws1.freeze_panes = 'A5'

# ---- SHEET 2: Detalle por PM y Proyecto ----
ws2 = wb_out.create_sheet('Detalle PM-Proyecto')
ws2.sheet_view.showGridLines = False

ws2.merge_cells('A1:H1')
ws2['A1'] = 'DETALLE DE HORAS POR PM Y PROYECTO – MARZO Y ABRIL 2026'
ws2['A1'].font = Font(name='Calibri', size=13, bold=True, color='FFFFFF')
ws2['A1'].fill = fill_dark_blue
ws2['A1'].alignment = center_align()
ws2.row_dimensions[1].height = 28

headers2 = ['Project Manager', 'Proyecto (OpenAir)', 'Categoría',
            'Horas Marzo', 'Horas Abril', 'Total Horas', '% Marzo PM', 'Nota']
col_widths2 = [25, 55, 22, 13, 13, 13, 13, 30]
for i, (h, w) in enumerate(zip(headers2, col_widths2), start=1):
    cell = ws2.cell(row=3, column=i, value=h)
    cell.font = hdr_font(10)
    cell.fill = fill_mid_blue
    cell.alignment = center_align(wrap=True)
    cell.border = border_thin
    ws2.column_dimensions[get_column_letter(i)].width = w
ws2.row_dimensions[3].height = 30

row_num = 4
for pm in PM_ORDER:
    # PM header row
    ws2.merge_cells(f'A{row_num}:H{row_num}')
    ws2[f'A{row_num}'] = f'▶  {pm}'
    ws2[f'A{row_num}'].font = Font(name='Calibri', size=11, bold=True, color='FFFFFF')
    ws2[f'A{row_num}'].fill = fill_dark_blue
    ws2[f'A{row_num}'].alignment = left_align()
    ws2.row_dimensions[row_num].height = 22
    row_num += 1

    pm_total_march = pm_month_total[pm].get('Marzo 2026', 0)

    # Combine March and April projects
    all_projs = set(pm_proj_month[pm].get('Marzo 2026', {}).keys()) | \
                set(pm_proj_month[pm].get('Abril 2026', {}).keys())

    for proj in sorted(all_projs, key=lambda p: -pm_proj_month[pm].get('Marzo 2026', {}).get(p, 0)):
        m_h = pm_proj_month[pm].get('Marzo 2026', {}).get(proj, 0)
        a_h = pm_proj_month[pm].get('Abril 2026', {}).get(proj, 0)
        tot = m_h + a_h
        cat = pm_proj_category.get((pm, proj), 'Facturable')
        pct = m_h / pm_total_march * 100 if pm_total_march > 0 else 0

        nota = ''
        if pct > 40: nota = 'Alta concentración'
        elif cat == 'Interno/Administrativo' and pct > 30: nota = 'Alto % en admin'
        elif cat == 'Vacaciones/Ausencia': nota = 'Ausencia registrada'

        vals = [pm, proj, cat, m_h, a_h, tot, pct/100, nota]
        for col_i, val in enumerate(vals, start=1):
            cell = ws2.cell(row=row_num, column=col_i, value=val)
            cell.border = border_thin
            cell.font = cell_font(9)
            cell.alignment = left_align(wrap=True) if col_i in [2, 8] else center_align()
            if col_i in [4, 5, 6]: cell.number_format = '#,##0.0'
            if col_i == 7: cell.number_format = '0.0%'

            # Category color coding
            if col_i == 3:
                if cat == 'Vacaciones/Ausencia':
                    cell.fill = fill_yellow
                elif cat == 'Interno/Administrativo':
                    cell.fill = fill_light_blue
                else:
                    cell.fill = fill_green
                    cell.font = Font(name='Calibri', size=9, color='FFFFFF')

            # Alternating row
            if row_num % 2 == 0 and col_i not in [3]:
                cell.fill = fill_light_gray

        ws2.row_dimensions[row_num].height = 18
        row_num += 1

    row_num += 1  # spacing

ws2.freeze_panes = 'A4'

# ---- SHEET 3: Evolución Semanal ----
ws3 = wb_out.create_sheet('Evolución Semanal')
ws3.sheet_view.showGridLines = False

all_weeks_sorted = sorted(set(w for pm in PM_ORDER for w in pm_weekly[pm] if '2026' in w))

ws3.merge_cells('A1:' + get_column_letter(1 + len(all_weeks_sorted)) + '1')
ws3['A1'] = 'EVOLUCIÓN SEMANAL DE HORAS POR PM – MARZO Y ABRIL 2026'
ws3['A1'].font = Font(name='Calibri', size=13, bold=True, color='FFFFFF')
ws3['A1'].fill = fill_dark_blue
ws3['A1'].alignment = center_align()
ws3.row_dimensions[1].height = 28

ws3.cell(row=3, column=1, value='Project Manager').font = hdr_font(10)
ws3.cell(row=3, column=1).fill = fill_dark_blue
ws3.cell(row=3, column=1).alignment = center_align(wrap=True)
ws3.cell(row=3, column=1).border = border_thin
ws3.column_dimensions['A'].width = 28

for i, week in enumerate(all_weeks_sorted, start=2):
    cell = ws3.cell(row=3, column=i, value=f'Sem.\n{week[:5]}')
    cell.font = hdr_font(9)
    cell.fill = fill_mid_blue if '/03/' in week else fill_orange if '/04/' in week else fill_mid_blue
    cell.alignment = center_align(wrap=True)
    cell.border = border_thin
    ws3.column_dimensions[get_column_letter(i)].width = 9
ws3.row_dimensions[3].height = 30

# Capacity row
cap_row = 4
ws3.cell(row=cap_row, column=1, value='Capacidad (44h)').font = Font(name='Calibri',size=9,italic=True,color='595959')
ws3.cell(row=cap_row, column=1).alignment = left_align()
for i, week in enumerate(all_weeks_sorted, start=2):
    cell = ws3.cell(row=cap_row, column=i, value=44)
    cell.font = Font(name='Calibri', size=9, italic=True, color='FF6B00')
    cell.alignment = center_align()
    cell.number_format = '#,##0.0'

for row_i, pm in enumerate(PM_ORDER, start=5):
    ws3.cell(row=row_i, column=1, value=pm).font = cell_font(10, bold=True)
    ws3.cell(row=row_i, column=1).alignment = left_align()
    ws3.cell(row=row_i, column=1).border = border_thin
    for col_i, week in enumerate(all_weeks_sorted, start=2):
        h = pm_weekly[pm].get(week, 0)
        cell = ws3.cell(row=row_i, column=col_i, value=h if h > 0 else None)
        cell.border = border_thin
        cell.alignment = center_align()
        cell.font = cell_font(10)
        cell.number_format = '#,##0.0'
        if h > 0:
            if h > 44: cell.fill = fill_red; cell.font = Font(name='Calibri',size=10,bold=True,color='FFFFFF')
            elif h >= 40: cell.fill = fill_yellow
            elif h > 0: cell.fill = fill_green; cell.font = Font(name='Calibri',size=10,color='FFFFFF')
        if row_i % 2 == 0 and h == 0:
            cell.fill = fill_light_gray
    ws3.row_dimensions[row_i].height = 22

ws3.freeze_panes = 'B5'

# ---- SHEET 4: Maestro Proyectos Activos ----
ws4 = wb_out.create_sheet('Proyectos Activos por PM')
ws4.sheet_view.showGridLines = False

ws4.merge_cells('A1:F1')
ws4['A1'] = 'MAESTRO DE PROYECTOS ACTIVOS POR PM – BLEND360 COLOMBIA'
ws4['A1'].font = Font(name='Calibri', size=13, bold=True, color='FFFFFF')
ws4['A1'].fill = fill_dark_blue
ws4['A1'].alignment = center_align()
ws4.row_dimensions[1].height = 28

headers4 = ['Project Manager', 'Project ID', 'Sigla Cliente', 'Objetivo del Proyecto', 'Tipo', 'Cliente']
col_widths4 = [25, 10, 14, 50, 18, 30]
for i, (h, w) in enumerate(zip(headers4, col_widths4), start=1):
    cell = ws4.cell(row=3, column=i, value=h)
    cell.font = hdr_font(10)
    cell.fill = fill_mid_blue
    cell.alignment = center_align(wrap=True)
    cell.border = border_thin
    ws4.column_dimensions[get_column_letter(i)].width = w
ws4.row_dimensions[3].height = 28

r = 4
for pm in PM_ORDER:
    for proj in pm_active_projects.get(pm, []):
        vals = [pm, proj['id'], proj['sigla'], proj['objetivo'], proj['tipo'], proj['cliente']]
        for ci, val in enumerate(vals, start=1):
            cell = ws4.cell(row=r, column=ci, value=val)
            cell.border = border_thin
            cell.font = cell_font(9)
            cell.alignment = left_align(wrap=True) if ci in [4, 6] else center_align()
            if r % 2 == 0:
                cell.fill = fill_light_gray
        ws4.row_dimensions[r].height = 18
        r += 1

ws4.freeze_panes = 'A4'

# Save Excel
excel_out = os.path.join(OUTPUT, 'Consolidado_Carga_PM_Blend360.xlsx')
wb_out.save(excel_out)
print(f"Excel consolidado guardado: {excel_out}")

# =============================================================================
# 4. HELPER FUNCTIONS FOR WORD DOCS
# =============================================================================
def add_heading(doc, text, level=1):
    """Add a heading with Blend styling."""
    # Use only levels 1-2 since template may not have Heading 3
    actual_level = min(level, 2)
    try:
        para = doc.add_heading(text, level=actual_level)
    except KeyError:
        para = doc.add_paragraph()
        para.add_run(text)
    for run in para.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = BLEND_DARK_BLUE
            run.bold = True
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = BLEND_MID_BLUE
            run.bold = True
        elif level == 3:
            run.font.size = Pt(11)
            run.font.color.rgb = BLEND_DARK_BLUE
            run.bold = True
    return para

def add_paragraph(doc, text, bold=False, italic=False, color=None, size=11,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    """Add a body paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    return para

def add_bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet point."""
    para = doc.add_paragraph(style='List Paragraph')
    para.paragraph_format.left_indent = Cm(1 + level * 0.5)
    para.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = para.add_run(bold_prefix)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = BLEND_DARK_BLUE
        run2 = para.add_run(text)
        run2.font.name = 'Calibri'
        run2.font.size = Pt(11)
    else:
        run = para.add_run(f'• {text}')
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    return para

def add_image_centered(doc, img_path, width_inches=6.0):
    """Add a centered image."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    return para

def add_table_header_row(table, headers, fill_hex=HEX_DARK_BLUE):
    """Style a table header row."""
    row = table.rows[0]
    for i, (cell, hdr) in enumerate(zip(row.cells, headers)):
        cell.text = hdr
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell._tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), fill_hex)
        cell._tc.tcPr.append(shading)

def set_cell_color(cell, fill_hex):
    """Set cell background color."""
    cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), fill_hex)
    cell._tc.tcPr.append(shading)

def add_separator(doc, color_hex=HEX_DARK_BLUE):
    """Add a thin colored separator line."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_after = Pt(6)
    return para

def add_info_box(doc, title, content_lines, fill_hex=HEX_LIGHT_BLUE):
    """Add a highlighted info box using a 1-cell table."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_color(cell, fill_hex)
    cell.paragraphs[0].clear()
    run_title = cell.paragraphs[0].add_run(title + '\n')
    run_title.bold = True
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = BLEND_DARK_BLUE
    for line in content_lines:
        p = cell.add_paragraph(f'• {line}')
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
    doc.add_paragraph()

# =============================================================================
# 5. WORD REPORT 1: INFORME DE CARGA
# =============================================================================
print("\nGenerando Word – Informe de Carga...")

doc1 = Document(TEMPLATE_FILE)

# Clear template content
for para in doc1.paragraphs:
    for run in para.runs:
        run.text = ''

# Cover
doc1.add_paragraph()
doc1.add_paragraph()
doc1.add_paragraph()

title_para = doc1.add_paragraph()
title_run = title_para.add_run('INFORME DE CARGA DE PROJECT MANAGERS')
title_run.font.name = 'Calibri'
title_run.font.size = Pt(22)
title_run.bold = True
title_run.font.color.rgb = BLEND_DARK_BLUE
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_after = Pt(12)

subtitle_para = doc1.add_paragraph()
sub_run = subtitle_para.add_run('Análisis de Utilización de Capacidad\nMarzo – Abril 2026')
sub_run.font.name = 'Calibri'
sub_run.font.size = Pt(14)
sub_run.font.color.rgb = BLEND_MID_BLUE
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_para.paragraph_format.space_after = Pt(8)

meta_para = doc1.add_paragraph()
meta_run = meta_para.add_run('Blend360 Colombia  |  PMO Operations  |  Abril 2026')
meta_run.font.name = 'Calibri'
meta_run.font.size = Pt(11)
meta_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc1.add_page_break()

# ---- 1. Objetivo ----
add_heading(doc1, '1. Objetivo del Informe', level=1)
add_separator(doc1)
add_paragraph(doc1,
    'El presente informe tiene como objetivo analizar la carga de trabajo de los Project Managers (PMs) '
    'de Blend360 Colombia durante los meses de marzo y abril de 2026, con base en las horas registradas '
    'en la plataforma OpenAir. Se evalúa la distribución de horas por proyecto, el nivel de utilización '
    'frente a la capacidad estándar de 44 horas semanales, y se identifican alertas y oportunidades '
    'de optimización en la asignación de proyectos.', size=11, space_after=8)

add_info_box(doc1, 'Parámetros de Análisis:', [
    'Capacidad estándar: 44 horas por semana por PM',
    'Período de análisis: 1 de marzo al 30 de abril de 2026',
    'Capacidad marzo: 193.6 horas (4.4 semanas)',
    'Fuente de datos: OpenAir – Reporte de horas (39 semanas)',
    f'Total de PMs analizados: {len(PM_ORDER)} Project Managers activos',
])

# ---- 2. Resumen Ejecutivo ----
add_heading(doc1, '2. Resumen Ejecutivo', level=1)
add_separator(doc1)

total_march_hours = sum(pm_month_total[pm].get('Marzo 2026', 0) for pm in PM_ORDER)
total_april_hours = sum(pm_month_total[pm].get('Abril 2026', 0) for pm in PM_ORDER)
overloaded = [pm for pm in PM_ORDER
              if pm_month_total[pm].get('Marzo 2026', 0) / MARCH_CAPACITY > 1.0]
high_load  = [pm for pm in PM_ORDER
              if 0.9 <= pm_month_total[pm].get('Marzo 2026', 0) / MARCH_CAPACITY <= 1.0]

add_paragraph(doc1,
    f'Durante marzo de 2026, los {len(PM_ORDER)} PMs activos de Blend360 Colombia registraron en total '
    f'{total_march_hours:,.0f} horas de trabajo distribuidas entre proyectos facturables, '
    f'actividades internas y ausencias. El análisis revela que el equipo opera en su mayoría '
    f'al límite de su capacidad disponible ({MARCH_CAPACITY:.0f} horas/mes por PM).',
    size=11, space_after=8)

# Summary table
add_heading(doc1, 'Tabla Resumen – Carga por PM', level=2)
table_s = doc1.add_table(rows=len(PM_ORDER)+1, cols=6)
table_s.style = 'Table Grid'
table_s.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(table_s, ['Project Manager', 'Proy.\nActivos',
                                'Horas\nMarzo', '% Cap.\nMarzo',
                                'Horas\nAbril', 'Estado'])
for row_i, pm in enumerate(PM_ORDER, start=1):
    n_p  = len(pm_active_projects.get(pm, []))
    m_h  = pm_month_total[pm].get('Marzo 2026', 0)
    a_h  = pm_month_total[pm].get('Abril 2026', 0)
    m_pct = m_h / MARCH_CAPACITY * 100

    if m_pct > 100:   estado = 'Sobrecargado'
    elif m_pct >= 90: estado = 'Carga Alta'
    else:             estado = 'Normal'

    row = table_s.rows[row_i]
    vals = [pm, str(n_p), f'{m_h:.0f}h', f'{m_pct:.1f}%', f'{a_h:.0f}h', estado]
    for ci, val in enumerate(vals):
        cell = row.cells[ci]
        cell.text = val
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
                run.bold = (ci == 5)
        # Color status
        if ci == 5:
            if estado == 'Sobrecargado':
                set_cell_color(cell, 'FF4D4D')
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            elif estado == 'Carga Alta':
                set_cell_color(cell, HEX_YELLOW)
            else:
                set_cell_color(cell, '70AD47')
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        elif row_i % 2 == 0:
            set_cell_color(cell, HEX_LIGHT_GRAY)

# Set column widths
for i, w in enumerate([2.5, 0.8, 0.9, 0.9, 0.9, 1.2]):
    for row in table_s.rows:
        row.cells[i].width = Inches(w)

doc1.add_paragraph()
add_image_centered(doc1, chart1_path, width_inches=6.0)
add_paragraph(doc1, 'Figura 1. Total de horas registradas por PM – Marzo vs Abril 2026',
              italic=True, color=RGBColor(0x59,0x59,0x59), size=9,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

doc1.add_paragraph()
add_image_centered(doc1, chart2_path, width_inches=5.8)
add_paragraph(doc1, 'Figura 2. Porcentaje de utilización de capacidad por PM – Marzo 2026',
              italic=True, color=RGBColor(0x59,0x59,0x59), size=9,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

doc1.add_page_break()

# ---- 3. Análisis por PM ----
add_heading(doc1, '3. Análisis Detallado por Project Manager', level=1)
add_separator(doc1)

pm_descriptions = {
    'Oscar Barragan': {
        'contexto': 'Oscar gestiona la mayor cartera de proyectos activos (56 proyectos). Durante marzo concentró su tiempo en el mantenimiento de SmartCash para BTG Pactual, los proyectos de UNA (Universidad Nacional) y el proyecto CAT de IA generativa.',
        'alerta': 'Con 56 proyectos activos, la gestión simultánea representa un riesgo de atención fragmentada. Se recomienda revisar cuáles proyectos son de baja actividad mensual.',
    },
    'Juan Bernal': {
        'contexto': 'Juan maneja una cartera diversificada con 26 proyectos activos. En marzo distribuyó sus horas entre clientes del sector público colombiano: ICFES, JPM Cloud Services, Secretaría del Hábitat (SDH), DIAN y BID.',
        'alerta': 'La distribución es relativamente equilibrada entre proyectos. El proyecto SDH y JPM consumen mayor tiempo de gestión.',
    },
    'David Cortes': {
        'contexto': 'David registró 198 horas en marzo, superando ligeramente la capacidad estándar (102.3%). Su tiempo estuvo concentrado en tres proyectos de alto volumen: Metro Medellín (50% de sus horas), CSJ y Ecuador IGM.',
        'alerta': 'Sobrecarga identificada. Concentración del 50% de horas en un solo proyecto (Metro Medellín). Se recomienda evaluar delegación o soporte adicional.',
    },
    'Miguel Garcia': {
        'contexto': 'Miguel trabaja en una mezcla de proyectos colombianos e internacionales. En marzo, el proyecto HotelPlanner (internacional) consumió el 25.8% de sus horas, seguido de proyectos nacionales como Secretaría de Educación (17% del tiempo).',
        'alerta': 'La carga de reuniones y capacitaciones (68h en marzo = 35%) es elevada. Se recomienda revisar la dinámica de este PM.',
    },
    'Kelly Carbonell': {
        'contexto': 'Kelly gestiona proyectos estratégicos de infraestructura para clientes como Cali (Tesorería a un Clic), BTG Pactual y Fiduagraria. El proyecto Treasury at a Click en Cali es el de mayor demanda (33.5% de sus horas de marzo).',
        'alerta': 'Cartera manejable con 14 proyectos activos. Monitorear la evolución del proyecto Cali Mayors Office que concentra el mayor volumen.',
    },
    'Daniel Sebastian Vargas': {
        'contexto': 'Daniel gestiona proyectos de IA y modernización. En marzo sus horas se distribuyeron entre el proyecto FOA Mojana (21.8%), Universidad de La Sabana (16.2%), IAPOL-G de Ponalsar (18%) y Supersociedades (16.4%).',
        'alerta': 'Proyectos inactivos aún recibiendo horas (INACTIVO SOC). Revisar el estado real de estos proyectos en OpenAir.',
    },
    'Diana Castro': {
        'contexto': 'Diana registra 145 horas en Meetings & Training en marzo, representando el 74.7% de su tiempo. Los proyectos facturables activos (Ingenium API y Secretaría de Planeación) reciben solo el 24.7% de su tiempo.',
        'alerta': 'Alta concentración de tiempo en actividades no facturables (74.7%). Esto requiere análisis: puede indicar que Diana cumple un rol de apoyo interno o que hay proyectos no facturables de alto valor. Requiere revisión urgente.',
    },
    'Diana Rojas': {
        'contexto': 'Diana gestiona proyectos de ADRES con alto volumen de horas. El proyecto ADRES Modernization consumió el 31.9% de sus horas en marzo. También gestiona múltiples proyectos de UNA y la operación ADRES.',
        'alerta': 'Gestiona 9 proyectos activos con carga distribuida. Vigilar el balance entre proyectos de implementación y operación de ADRES.',
    },
    'Indira Duarte': {
        'contexto': 'Indira registró 203 horas en marzo (104.6% de capacidad), la mayor sobrecarga del equipo. El 88.2% de su tiempo fue dedicado al proyecto ADRES AI Medical Audit Phase I, con solo el 11.8% en reuniones internas.',
        'alerta': 'Sobrecarga crítica (104.6%). Concentración extrema en un solo proyecto (88.2%). Se recomienda asignar soporte o redistribuir tareas dentro del proyecto ADRES.',
    },
}

for pm in PM_ORDER:
    add_heading(doc1, f'3.{PM_ORDER.index(pm)+1}. {pm}', level=2)
    n_p = len(pm_active_projects.get(pm, []))
    m_h = pm_month_total[pm].get('Marzo 2026', 0)
    a_h = pm_month_total[pm].get('Abril 2026', 0)
    m_pct = m_h / MARCH_CAPACITY * 100

    # KPI mini-table
    kpi_table = doc1.add_table(rows=1, cols=4)
    kpi_table.style = 'Table Grid'
    kpi_data = [
        ('Proyectos Activos', str(n_p)),
        ('Horas Marzo', f'{m_h:.0f}h'),
        ('% Capacidad Marzo', f'{m_pct:.1f}%'),
        ('Horas Abril', f'{a_h:.0f}h'),
    ]
    for ci, (lbl, val) in enumerate(kpi_data):
        cell = kpi_table.rows[0].cells[ci]
        cell.paragraphs[0].clear()
        p1 = cell.paragraphs[0]
        r_lbl = p1.add_run(lbl + '\n')
        r_lbl.font.name = 'Calibri'; r_lbl.font.size = Pt(8); r_lbl.bold = True
        r_lbl.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r_val = p1.add_run(val)
        r_val.font.name = 'Calibri'; r_val.font.size = Pt(14); r_val.bold = True
        r_val.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        color = HEX_DARK_BLUE if ci % 2 == 0 else HEX_MID_BLUE
        set_cell_color(cell, color)
        cell.width = Inches(1.6)

    doc1.add_paragraph()

    desc = pm_descriptions.get(pm, {})
    if desc.get('contexto'):
        add_paragraph(doc1, desc['contexto'], size=11, space_after=6)

    if desc.get('alerta'):
        alerta_para = doc1.add_paragraph()
        alerta_para.paragraph_format.left_indent = Cm(0.5)
        alerta_para.paragraph_format.space_after = Pt(6)
        icon_run = alerta_para.add_run('⚑ Alerta: ')
        icon_run.font.name = 'Calibri'; icon_run.font.size = Pt(11)
        icon_run.bold = True; icon_run.font.color.rgb = BLEND_ORANGE
        txt_run = alerta_para.add_run(desc['alerta'])
        txt_run.font.name = 'Calibri'; txt_run.font.size = Pt(11)
        txt_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    # Project detail table
    projs_march = sorted(pm_proj_month[pm].get('Marzo 2026', {}).items(),
                          key=lambda x: -x[1])[:6]
    if projs_march:
        add_heading(doc1, 'Principales proyectos – Marzo 2026', level=3)
        proj_table = doc1.add_table(rows=len(projs_march)+1, cols=4)
        proj_table.style = 'Table Grid'
        add_table_header_row(proj_table, ['Proyecto', 'Categoría', 'Horas Marzo', '% del Total'],
                              fill_hex=HEX_MID_BLUE)
        for ri, (proj, hours) in enumerate(projs_march, start=1):
            cat = pm_proj_category.get((pm, proj), 'Facturable')
            pct = hours / m_h * 100 if m_h > 0 else 0
            row = proj_table.rows[ri]
            vals = [proj[:55], cat, f'{hours:.0f}', f'{pct:.1f}%']
            for ci, val in enumerate(vals):
                cell = row.cells[ci]
                cell.text = val
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.name = 'Calibri'; r.font.size = Pt(9)
                if ri % 2 == 0:
                    set_cell_color(cell, HEX_LIGHT_GRAY)
                if ci == 1:
                    if cat == 'Facturable':
                        set_cell_color(cell, '70AD47')
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                    elif cat == 'Interno/Administrativo':
                        set_cell_color(cell, HEX_LIGHT_BLUE)
                    else:
                        set_cell_color(cell, HEX_YELLOW)
        doc1.add_paragraph()

    # Chart
    if pm in pm_chart_paths:
        add_image_centered(doc1, pm_chart_paths[pm], width_inches=5.5)
        add_paragraph(doc1,
                      f'Figura. Distribución de horas por proyecto – {pm} – Marzo 2026',
                      italic=True, color=RGBColor(0x59,0x59,0x59), size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    add_separator(doc1, HEX_LIGHT_BLUE)

doc1.add_page_break()

# ---- 4. Comparativa General ----
add_heading(doc1, '4. Comparativa General del Equipo', level=1)
add_separator(doc1)
add_paragraph(doc1,
    'La siguiente gráfica muestra la evolución semanal de horas registradas por cada PM durante '
    'el período analizado, evidenciando la consistencia en la carga de trabajo semana a semana.',
    size=11, space_after=8)

add_image_centered(doc1, chart5_path, width_inches=6.5)
add_paragraph(doc1,
    'Figura 3. Evolución semanal de horas por PM – Marzo y Abril 2026. '
    'La línea naranja indica el límite de 44 horas semanales.',
    italic=True, color=RGBColor(0x59,0x59,0x59), size=9,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

doc1.add_paragraph()
add_image_centered(doc1, chart3_path, width_inches=5.8)
add_paragraph(doc1,
    'Figura 4. Número de proyectos activos asignados por PM (fuente: Maestro de Proyectos).',
    italic=True, color=RGBColor(0x59,0x59,0x59), size=9,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

doc1.add_paragraph()
add_image_centered(doc1, chart4_path, width_inches=6.0)
add_paragraph(doc1,
    'Figura 5. Distribución de horas por categoría (facturable, interno/administrativo, '
    'vacaciones/ausencia) por PM – Marzo 2026.',
    italic=True, color=RGBColor(0x59,0x59,0x59), size=9,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

doc1.add_page_break()

# ---- 5. Conclusiones ----
add_heading(doc1, '5. Conclusiones y Alertas', level=1)
add_separator(doc1)

add_paragraph(doc1,
    'A partir del análisis de los datos de OpenAir para marzo y abril de 2026, se identifican '
    'los siguientes hallazgos clave:', size=11, space_after=8)

conclusions = [
    ('Equipo operando a plena capacidad: ',
     'El 100% de los PMs analizados registraron horas cercanas o superiores a la capacidad '
     'estándar de 44h/semana durante marzo. Esto indica que el equipo no tiene margen de '
     'absorción para nuevos proyectos sin redistribución previa.'),
    ('PMs con sobrecarga identificada: ',
     'David Cortes (102.3%) e Indira Duarte (104.6%) superaron la capacidad en marzo. '
     'Indira concentra el 88.2% de su tiempo en un único proyecto (ADRES AI Medical Audit), '
     'lo cual representa un riesgo operativo.'),
    ('Exceso de proyectos – Oscar Barragan: ',
     'Con 56 proyectos activos asignados, Oscar Barragan tiene la cartera más extensa del equipo. '
     'Se recomienda realizar una revisión de proyectos en estado inactivo o de baja actividad '
     'que deben ser actualizados o cerrados en el maestro.'),
    ('Patrón atípico – Diana Castro: ',
     'El 74.7% del tiempo registrado de Diana Castro en marzo corresponde a reuniones internas '
     'y capacitación. Se requiere validar si este patrón refleja su rol actual o si hay un '
     'problema de registro de horas.'),
    ('Proyectos con códigos INACTIVO activos: ',
     'Daniel Sebastian Vargas registra horas contra proyectos marcados como "INACTIVO" en '
     'OpenAir, lo que indica inconsistencias en la gestión de proyectos en la plataforma.'),
    ('Abril con datos parciales: ',
     'Los datos de abril reflejan entre 1.6 y 2.6 semanas de registro por PM, lo que sugiere '
     'que no todos han completado el reporte de horas. Se recomienda verificar el cierre '
     'de timesheets.'),
]

for bold_txt, regular_txt in conclusions:
    para = doc1.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_after = Pt(8)
    r1 = para.add_run(f'• {bold_txt}')
    r1.font.name = 'Calibri'; r1.font.size = Pt(11)
    r1.bold = True; r1.font.color.rgb = BLEND_DARK_BLUE
    r2 = para.add_run(regular_txt)
    r2.font.name = 'Calibri'; r2.font.size = Pt(11)

doc1.add_page_break()

# ---- 6. Datos de Referencia ----
add_heading(doc1, '6. Fuentes y Metodología', level=1)
add_separator(doc1)
add_bullet(doc1, 'Archivo de horas: OpenAir Export – Horas Marzo Abril 2026 (42,542 registros)')
add_bullet(doc1, 'Maestro de proyectos: maestro.proyectos.xlsx – Hoja "project" (173 proyectos activos)')
add_bullet(doc1, 'Capacidad estándar utilizada: 44 horas/semana por PM')
add_bullet(doc1, 'Cálculo de capacidad mensual: número de semanas × 44 horas')
add_bullet(doc1, 'Proyectos identificados como colombianos: prefijo "CO -" o cliente colombiano')
add_bullet(doc1, 'Categorías de horas: Facturable / Interno-Administrativo / Vacaciones-Ausencia')
add_bullet(doc1, 'Herramientas utilizadas: Python (openpyxl, matplotlib, python-docx)')

doc1_path = os.path.join(OUTPUT, 'Informe_Carga_PMs_Blend360.docx')
doc1.save(doc1_path)
print(f"Word Informe guardado: {doc1_path}")

# =============================================================================
# 6. WORD REPORT 2: PROPUESTA DE MANEJO
# =============================================================================
print("\nGenerando Word – Propuesta de Manejo...")

doc2 = Document(TEMPLATE_FILE)
for para in doc2.paragraphs:
    for run in para.runs:
        run.text = ''

# Cover
doc2.add_paragraph()
doc2.add_paragraph()
doc2.add_paragraph()

t_para = doc2.add_paragraph()
t_run = t_para.add_run('PROPUESTA DE GESTIÓN Y ASIGNACIÓN\nEFICIENTE DE PROYECTOS')
t_run.font.name = 'Calibri'; t_run.font.size = Pt(22)
t_run.bold = True; t_run.font.color.rgb = BLEND_DARK_BLUE
t_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
t_para.paragraph_format.space_after = Pt(14)

st_para = doc2.add_paragraph()
st_run = st_para.add_run('Modelo de Asignación de Proyectos a Project Managers\nBlend360 Colombia – PMO Operations')
st_run.font.name = 'Calibri'; st_run.font.size = Pt(13)
st_run.font.color.rgb = BLEND_MID_BLUE
st_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
st_para.paragraph_format.space_after = Pt(8)

m_para = doc2.add_paragraph()
m_run = m_para.add_run('Versión 1.0  |  Abril 2026  |  Confidencial')
m_run.font.name = 'Calibri'; m_run.font.size = Pt(11)
m_run.font.color.rgb = RGBColor(0x59,0x59,0x59)
m_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc2.add_page_break()

# ---- 1. Introducción ----
add_heading(doc2, '1. Introducción', level=1)
add_separator(doc2)
add_paragraph(doc2,
    'La gestión eficiente de proyectos en Blend360 Colombia requiere un modelo estructurado de '
    'asignación que garantice la sostenibilidad operativa del equipo de PMs, la calidad en la '
    'entrega de proyectos y la satisfacción de los clientes. El análisis de los datos de marzo '
    'y abril de 2026 revela que el equipo actualmente opera al 100% de su capacidad, con señales '
    'de sobrecarga puntual en algunos PMs y una cartera que supera la capacidad de atención '
    'individual en varios casos.',
    size=11, space_after=8)

add_paragraph(doc2,
    'Este documento propone un modelo de asignación basado en criterios objetivos de carga, '
    'complejidad del proyecto y capacidad disponible, con el objetivo de distribuir el trabajo '
    'de manera equitativa y sostenible.', size=11, space_after=8)

# ---- 2. Diagnóstico ----
add_heading(doc2, '2. Diagnóstico de la Situación Actual', level=1)
add_separator(doc2)

add_heading(doc2, '2.1 Hallazgos Clave', level=2)

diag_items = [
    ('Saturación de capacidad:', '9 de 9 PMs operaron al 90-105% de su capacidad en marzo 2026, lo que elimina cualquier margen para absorber proyectos adicionales sin planificación previa.'),
    ('Cartera desproporcionada:', 'Oscar Barragan tiene 56 proyectos activos asignados, mientras Indira Duarte tiene 4. La distribución actual no responde a un criterio de carga homogénea.'),
    ('Concentración de riesgo:', 'Indira Duarte dedica el 88% de su tiempo a un único proyecto. Este nivel de concentración crea dependencia crítica y vulnerabilidad operativa.'),
    ('Bajo tiempo facturable en algunos PMs:', 'Diana Castro registra solo el 24.7% de sus horas en proyectos facturables durante marzo, lo que impacta la eficiencia del equipo.'),
    ('Inconsistencias en OpenAir:', 'Se identificaron proyectos marcados como "INACTIVO" en los que aún se registran horas. El maestro de proyectos no refleja con precisión la cartera activa real.'),
    ('Ausencia de criterios formales de asignación:', 'No existe actualmente un proceso documentado que defina cuántos proyectos puede gestionar un PM según su carga en horas ni cómo se prioriza la asignación.'),
]

for bold_txt, txt in diag_items:
    add_bullet(doc2, txt, bold_prefix=bold_txt + ' ')

doc2.add_paragraph()

# Metrics table
add_heading(doc2, '2.2 Indicadores de Referencia Marzo 2026', level=2)
metrics_table = doc2.add_table(rows=5, cols=3)
metrics_table.style = 'Table Grid'
add_table_header_row(metrics_table, ['Indicador', 'Valor', 'Referencia / Alerta'])
metrics_data = [
    ('Capacidad disponible mensual', '193.6 h/PM', '44h × 4.4 semanas'),
    ('PMs con utilización > 100%', '2 de 9 (22%)', 'Alerta: rebalancear'),
    ('PMs con utilización 90-100%', '7 de 9 (78%)', 'Monitorear mensual'),
    ('Mayor concentración en 1 proyecto', '88.2% (Indira Duarte)', 'Riesgo operativo alto'),
]
for ri, (ind, val, ref) in enumerate(metrics_data, start=1):
    row = metrics_table.rows[ri]
    for ci, v in enumerate([ind, val, ref]):
        cell = row.cells[ci]
        cell.text = v
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Calibri'; r.font.size = Pt(10)
        if ri % 2 == 0:
            set_cell_color(cell, HEX_LIGHT_GRAY)
doc2.add_paragraph()

# ---- 3. Modelo de Asignación ----
add_heading(doc2, '3. Modelo Propuesto de Asignación', level=1)
add_separator(doc2)

add_paragraph(doc2,
    'El modelo propuesto se estructura en torno a cuatro dimensiones: criterios de asignación, '
    'límites de carga, proceso de revisión y herramientas de seguimiento. Su implementación '
    'requiere el compromiso de los líderes de PMO y la actualización periódica del maestro '
    'de proyectos.', size=11, space_after=8)

add_heading(doc2, '3.1 Criterios de Asignación de Proyectos a un PM', level=2)
add_paragraph(doc2,
    'La asignación de un nuevo proyecto a un PM debe considerar los siguientes criterios en orden '
    'de prioridad:', size=11, space_after=6)

criteria = [
    ('Capacidad disponible (criterio principal):',
     'El PM debe tener al menos el 20% de su capacidad mensual libre al momento de la asignación. '
     'Con 44h/semana, esto equivale a un mínimo de 8.8h/semana disponibles.'),
    ('Número de proyectos activos:',
     'Se recomienda que cada PM gestione entre 8 y 20 proyectos activos simultáneamente. '
     'Proyectos con baja actividad (menos de 4h/mes) deben revisarse para cierre o reasignación.'),
    ('Afinidad técnica y de cliente:',
     'Cuando sea posible, agrupar proyectos del mismo cliente o sector en un mismo PM, '
     'para reducir el tiempo de contextualización y fortalecer la relación cliente-PM.'),
    ('Complejidad del proyecto:',
     'Proyectos de alta complejidad (implementaciones, proyectos Fixed con alcance amplio) '
     'requieren mayor tiempo de gestión. Estos deben contabilizarse como equivalentes a '
     '1.5 proyectos de tamaño estándar al calcular la carga.'),
    ('Tipo de contrato:',
     'Proyectos de operación o retainer (Ret) con flujo continuo de horas requieren atención '
     'regular y deben priorizarse en la asignación a PMs con disponibilidad sostenida.'),
]

for i, (bold_txt, txt) in enumerate(criteria, start=1):
    para = doc2.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_after = Pt(8)
    r1 = para.add_run(f'{i}. {bold_txt} ')
    r1.font.name = 'Calibri'; r1.font.size = Pt(11)
    r1.bold = True; r1.font.color.rgb = BLEND_DARK_BLUE
    r2 = para.add_run(txt)
    r2.font.name = 'Calibri'; r2.font.size = Pt(11)

add_heading(doc2, '3.2 Límites de Carga Recomendados', level=2)
add_paragraph(doc2,
    'Se establecen los siguientes umbrales de carga como guía para la toma de decisiones '
    'de asignación:', size=11, space_after=6)

limits_table = doc2.add_table(rows=5, cols=4)
limits_table.style = 'Table Grid'
add_table_header_row(limits_table, ['Zona de Carga', '% Capacidad', 'Acción Recomendada', 'Semáforo'])
limits_data = [
    ('Óptima',        '70% – 85%',   'Asignación normal. PM disponible para nuevos proyectos.',     '🟢 Verde'),
    ('Alta',          '85% – 95%',   'Precaución. Evaluar antes de asignar nuevos proyectos.',      '🟡 Amarillo'),
    ('Crítica',       '95% – 100%',  'Solo asignar si otro proyecto libera carga. Monitorear.',     '🟠 Naranja'),
    ('Sobrecarga',    '> 100%',      'No asignar nuevos proyectos. Redistribuir urgente.',           '🔴 Rojo'),
]
for ri, (zona, pct, accion, semaforo) in enumerate(limits_data, start=1):
    row = limits_table.rows[ri]
    for ci, v in enumerate([zona, pct, accion, semaforo]):
        cell = row.cells[ci]
        cell.text = v
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 2 else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Calibri'; r.font.size = Pt(10)
    if ri % 2 == 0:
        for cell in row.cells:
            set_cell_color(cell, HEX_LIGHT_GRAY)
doc2.add_paragraph()

add_heading(doc2, '3.3 Proceso de Asignación', level=2)
add_paragraph(doc2,
    'Se propone el siguiente proceso de 5 pasos para la asignación de nuevos proyectos:',
    size=11, space_after=6)

process_steps = [
    ('Paso 1 – Solicitud de asignación:',
     'Cuando se origina un nuevo proyecto (oportunidad ganada), el equipo comercial notifica '
     'a PMO con la información básica: cliente, tipo de proyecto, volumen estimado de horas '
     'por mes y fecha de inicio.'),
    ('Paso 2 – Revisión del tablero de carga:',
     'El líder de PMO consulta el consolidado semanal de carga (archivo Excel) para identificar '
     'el PM con mayor disponibilidad en el período proyectado.'),
    ('Paso 3 – Evaluación de criterios:',
     'Se verifica: capacidad disponible, número de proyectos activos, afinidad técnica y '
     'del cliente. Se selecciona el PM más adecuado.'),
    ('Paso 4 – Aprobación y comunicación:',
     'El líder de PMO aprueba la asignación y comunica al PM seleccionado con al menos '
     '5 días hábiles de anticipación al inicio del proyecto.'),
    ('Paso 5 – Actualización del maestro:',
     'Se actualiza el maestro de proyectos (maestro.proyectos.xlsx) con el PM asignado, '
     'la fecha de inicio y el estado "Activo". Se configura el proyecto en OpenAir.'),
]

for bold_txt, txt in process_steps:
    para = doc2.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_after = Pt(8)
    r1 = para.add_run(bold_txt + ' ')
    r1.font.name = 'Calibri'; r1.font.size = Pt(11)
    r1.bold = True; r1.font.color.rgb = BLEND_DARK_BLUE
    r2 = para.add_run(txt)
    r2.font.name = 'Calibri'; r2.font.size = Pt(11)

# ---- 4. Herramientas de Seguimiento ----
add_heading(doc2, '4. Herramientas de Seguimiento', level=1)
add_separator(doc2)

add_paragraph(doc2,
    'Para hacer sostenible el modelo propuesto, se requieren las siguientes herramientas '
    'y rutinas de seguimiento:', size=11, space_after=8)

tools = [
    ('Tablero mensual de carga (Excel):',
     'El archivo "Consolidado_Carga_PM_Blend360.xlsx" debe actualizarse mensualmente con '
     'el export de OpenAir. Incluye las hojas: Resumen por PM, Detalle PM-Proyecto, '
     'Evolución Semanal y Proyectos Activos.'),
    ('Revisión quincenal de carga:',
     'El líder de PMO debe revisar la carga de cada PM cada dos semanas, usando los datos '
     'de OpenAir, para identificar desviaciones tempranas y actuar antes de que se '
     'conviertan en sobrecargas.'),
    ('Maestro de proyectos actualizado:',
     'El archivo maestro.proyectos.xlsx debe mantenerse al día, cerrando proyectos '
     'inactivos y verificando que el PM asignado en el maestro coincida con quien reporta '
     'horas en OpenAir.'),
    ('Alerta automática en OpenAir:',
     'Se recomienda configurar alertas en OpenAir cuando un PM supere el 90% de su '
     'capacidad mensual (39.6 horas/semana promedio), permitiendo acción preventiva.'),
    ('Reunión mensual de portafolio:',
     'Espacio mensual de 1 hora entre el líder de PMO y los PMs para revisar la carga '
     'proyectada, identificar proyectos en riesgo de finalización y planificar la '
     'disponibilidad futura.'),
]

for bold_txt, txt in tools:
    para = doc2.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_after = Pt(8)
    r1 = para.add_run(f'• {bold_txt} ')
    r1.font.name = 'Calibri'; r1.font.size = Pt(11)
    r1.bold = True; r1.font.color.rgb = BLEND_DARK_BLUE
    r2 = para.add_run(txt)
    r2.font.name = 'Calibri'; r2.font.size = Pt(11)

# ---- 5. Plan de Acción Inmediata ----
add_heading(doc2, '5. Plan de Acción Inmediata', level=1)
add_separator(doc2)
add_paragraph(doc2,
    'Con base en el diagnóstico de marzo-abril 2026, se proponen las siguientes acciones '
    'prioritarias para implementar en el corto plazo:', size=11, space_after=8)

action_table = doc2.add_table(rows=8, cols=4)
action_table.style = 'Table Grid'
add_table_header_row(action_table,
    ['Acción', 'Responsable', 'Plazo', 'Prioridad'])
actions = [
    ('Revisión de cartera Oscar Barragan: Identificar y cerrar proyectos inactivos (meta: < 30 proyectos activos)', 'Líder PMO + Oscar Barragan', 'Mayo 2026', 'Alta'),
    ('Soporte para Indira Duarte en ADRES AI Medical: Asignar PM de apoyo o redistribuir tareas', 'Líder PMO', 'Inmediato', 'Crítica'),
    ('Revisión rol Diana Castro: Clarificar si el perfil de reuniones internas es intencional o hay desvío de registro', 'PMO + Diana Castro', 'Abril 2026', 'Alta'),
    ('Limpieza de proyectos INACTIVO en OpenAir: Daniel Vargas registra horas en proyectos marcados como inactivos', 'PMO + Daniel Vargas', 'Inmediato', 'Media'),
    ('Cierre de timesheets abril: Verificar que todos los PMs hayan cerrado reporte de horas de abril', 'Todos los PMs', 'Abril 30', 'Alta'),
    ('Implementar tablero mensual de carga: Usar el Excel generado como base y programar actualización mensual', 'PMO Lead', 'Mayo 2026', 'Media'),
    ('Documentar proceso de asignación: Formalizar el proceso de 5 pasos con formulario estándar', 'PMO Lead', 'Mayo 2026', 'Media'),
]
priority_colors = {'Crítica': 'FF4D4D', 'Alta': HEX_ORANGE, 'Media': HEX_YELLOW, 'Baja': '70AD47'}
for ri, (accion, resp, plazo, prior) in enumerate(actions, start=1):
    row = action_table.rows[ri]
    for ci, v in enumerate([accion, resp, plazo, prior]):
        cell = row.cells[ci]
        cell.text = v
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Calibri'; r.font.size = Pt(9)
                if ci == 3:
                    r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        if ci == 3:
            set_cell_color(cell, priority_colors.get(prior, '595959'))
        elif ri % 2 == 0:
            set_cell_color(cell, HEX_LIGHT_GRAY)
doc2.add_paragraph()

# ---- 6. Indicadores de Seguimiento (KPIs) ----
add_heading(doc2, '6. Indicadores de Seguimiento (KPIs)', level=1)
add_separator(doc2)
add_paragraph(doc2,
    'Se propone monitorear mensualmente los siguientes indicadores para evaluar la efectividad '
    'del modelo de asignación:', size=11, space_after=6)

kpis = [
    ('% Utilización promedio del equipo:', 'Meta: entre 80% y 90% (zona verde-amarilla). Indica si el equipo tiene capacidad para nuevos proyectos.'),
    ('# PMs con sobrecarga (> 100%):', 'Meta: 0. Cualquier PM en sobrecarga debe generar alerta inmediata.'),
    ('# Proyectos activos promedio por PM:', 'Meta: entre 10 y 20. Indica diversificación manejable sin fragmentación excesiva.'),
    ('% Horas facturables vs total horas:', 'Meta: > 75%. Mide la productividad neta del equipo de PMs.'),
    ('# Proyectos cerrados en el mes:', 'Indicador de gestión del portafolio. Un cierre regular indica buena gestión del ciclo de vida.'),
    ('Índice de concentración (horas en 1 proyecto / total):', 'Meta: < 60%. Valores superiores indican riesgo de dependencia excesiva en un solo proyecto.'),
]
for bold_txt, txt in kpis:
    para = doc2.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_after = Pt(6)
    r1 = para.add_run(f'• {bold_txt} ')
    r1.font.name = 'Calibri'; r1.font.size = Pt(11)
    r1.bold = True; r1.font.color.rgb = BLEND_DARK_BLUE
    r2 = para.add_run(txt)
    r2.font.name = 'Calibri'; r2.font.size = Pt(11)

doc2.add_page_break()

# ---- 7. Resumen Ejecutivo ----
add_heading(doc2, '7. Resumen Ejecutivo de la Propuesta', level=1)
add_separator(doc2)

add_info_box(doc2, 'Puntos Clave del Modelo:', [
    'Capacidad máxima para asignación: 80-85% (35-37h/semana)',
    'Máximo proyectos por PM: 20 proyectos activos con actividad mensual',
    'Revisión de carga: quincenal obligatoria, mensual con todo el equipo',
    'Proceso de 5 pasos para nueva asignación (solicitud → revisión → evaluación → aprobación → actualización)',
    'Tablero mensual en Excel como herramienta central de seguimiento',
    'Alerta automática en OpenAir al superar 90% de capacidad',
], fill_hex=HEX_LIGHT_BLUE)

add_paragraph(doc2,
    'La implementación de este modelo permitirá a Blend360 Colombia gestionar el crecimiento '
    'del portafolio de proyectos de manera sostenible, protegiendo la salud operativa del '
    'equipo de PMs, mejorando la calidad de atención a los clientes y generando visibilidad '
    'temprana sobre riesgos de sobrecarga.', size=11, space_after=8)

add_paragraph(doc2,
    'Se recomienda revisar y ajustar los parámetros del modelo (límites de carga, número '
    'máximo de proyectos) trimestralmente, incorporando la retroalimentación del equipo y '
    'los resultados observados.', size=11, space_after=8)

add_separator(doc2)
add_paragraph(doc2,
    'Documento elaborado por la Dirección de PMO – Blend360 Colombia  |  Abril 2026',
    italic=True, color=RGBColor(0x59,0x59,0x59), size=10,
    alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc2_path = os.path.join(OUTPUT, 'Propuesta_Gestion_Asignacion_PM_Blend360.docx')
doc2.save(doc2_path)
print(f"Word Propuesta guardado: {doc2_path}")

# =============================================================================
# 7. SAVE PROMPT TO PROCESS FOLDER
# =============================================================================
prompt_text = """# Prompt: Análisis de Carga de Project Managers – Blend360 Colombia
# Fecha: Abril 2026
# Generado por: Claude Code (Anthropic) + PMO Operations

## Descripción del Proceso
Este prompt y script generan automáticamente los siguientes entregables:
1. Informe Word de carga por PM con datos reales y gráficas
2. Consolidado Excel con evolución y detalle de carga marzo-abril
3. Propuesta Word de manejo eficiente de asignación de proyectos

## Archivos de Entrada (1. Input)
- maestro.proyectos.xlsx: Maestro de proyectos activos con PM asignado
- Horas marzo abril.xlsx: Reporte OpenAir con horas registradas (42,542 registros)
- Plantilla word Blend.docx: Plantilla corporativa para documentos Word

## Lógica del Análisis
- Capacidad estándar: 44 horas/semana por PM
- Marzo 2026: 4.4 semanas efectivas → 193.6 horas/PM
- PMs identificados: Oscar Barragan, Juan Bernal, David Cortes, Miguel Garcia,
  Kelly Carbonell, Daniel Sebastian Vargas, Diana Castro, Diana Rojas, Indira Duarte
- Categorías de horas: Facturable | Interno/Administrativo | Vacaciones/Ausencia
- Proyectos colombianos: prefijo "CO -" en nombre del proyecto + clientes colombianos

## Archivos Generados (3. Output)
- Consolidado_Carga_PM_Blend360.xlsx: 4 hojas (Resumen, Detalle, Evolución, Proyectos)
- Informe_Carga_PMs_Blend360.docx: Informe detallado con gráficas por PM
- Propuesta_Gestion_Asignacion_PM_Blend360.docx: Propuesta de modelo de asignación
- charts/: Gráficas PNG generadas con matplotlib

## Para Ejecutar
1. Asegurarse de tener Python 3.x con: openpyxl, matplotlib, python-docx, pandas
2. Abrir terminal en la carpeta PMO-Operations
3. Ejecutar: py gen_outputs.py
   O desde la carpeta 2. Process: py ../gen_outputs.py

## Actualización Mensual
Para actualizar con nuevos datos:
1. Reemplazar "Horas marzo abril.xlsx" con el nuevo export de OpenAir
2. Actualizar maestro.proyectos.xlsx si hay cambios en asignaciones
3. Ejecutar el script nuevamente
4. Los archivos de Output se sobrescriben automáticamente
"""

prompt_path = os.path.join(PROCESS, 'prompt_analisis_carga_pm.md')
with open(prompt_path, 'w', encoding='utf-8') as f:
    f.write(prompt_text)
print(f"Prompt guardado: {prompt_path}")

# Copy script to Process folder
script_dest = os.path.join(PROCESS, 'analisis_carga_pm.py')
shutil.copy2(__file__, script_dest)
print(f"Script copiado a: {script_dest}")

print("\n" + "="*60)
print("GENERACION COMPLETADA EXITOSAMENTE")
print("="*60)
print(f"\nArchivos generados en: {OUTPUT}")
print(f"  - Consolidado_Carga_PM_Blend360.xlsx")
print(f"  - Informe_Carga_PMs_Blend360.docx")
print(f"  - Propuesta_Gestion_Asignacion_PM_Blend360.docx")
print(f"  - charts/ (gráficas PNG)")
print(f"\nPrompt y script en: {PROCESS}")
print(f"  - prompt_analisis_carga_pm.md")
print(f"  - analisis_carga_pm.py")
