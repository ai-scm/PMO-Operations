"""
Generador de Informe de Seguimiento - Centro de Servicios
Cruza: Daily 10/04/2026 + Tablero ClickUp P0321
Salida: Informe_Seguimiento_CentroServicios_20260414.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import copy
from pathlib import Path

BASE    = Path(r'c:\Users\BOG-LAP-SER-176\Documents\bld-engineering-pmo-col\Centro de Servicios')
INPUT   = BASE / '1. Input'
OUTPUT  = BASE / '3. Output'
OUTPUT.mkdir(exist_ok=True)

PLANTILLA = INPUT / 'Plantilla word Blend.docx'
SALIDA    = OUTPUT / 'Informe_Seguimiento_CentroServicios_20260414.docx'

# ─────────────────────────────────────────────
# Helpers de formato
# ─────────────────────────────────────────────
def set_titulo(doc, texto, nivel=1):
    """Montserrat 16 negrita"""
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1'] if nivel == 1 else doc.styles['Heading 2']
    run = p.add_run(texto)
    run.font.name = 'Montserrat'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def set_subtitulo(doc, texto):
    """Montserrat 14 negrita"""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name = 'Montserrat'
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def set_parrafo(doc, texto, justify=True):
    """Montserrat 10 sin negrita negro"""
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(texto)
    run.font.name = 'Montserrat'
    run.font.size = Pt(10)
    run.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def set_bullet(doc, texto):
    """Bullet Montserrat 10"""
    p = doc.add_paragraph(style='List Paragraph')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f'• {texto}')
    run.font.name = 'Montserrat'
    run.font.size = Pt(10)
    run.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_text(cell, texto, bold=False, size=10, color='000000'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(texto)
    run.font.name = 'Montserrat'
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(
        int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))

COLOR_URGENTE = 'FF4444'
COLOR_HIGH    = 'FF8C00'
COLOR_NORMAL  = '4472C4'
COLOR_LOW     = '70AD47'

PRIORIDAD_COLOR = {
    'URGENTE': 'FFCCCC',
    'ALTA':    'FFE5CC',
    'NORMAL':  'CCE0FF',
    'BAJA':    'D9F0CC',
}
PRIORIDAD_HEAT = {
    'URGENTE': 'C00000',
    'ALTA':    'BF5900',
    'NORMAL':  '1F3864',
    'BAJA':    '375623',
}

# ─────────────────────────────────────────────
# Crear documento desde plantilla
# ─────────────────────────────────────────────
doc = Document(PLANTILLA)

# Limpiar párrafos existentes del cuerpo (conservar header/footer)
for p in doc.paragraphs:
    p.clear()
for t in doc.tables:
    t._element.getparent().remove(t._element)

# Quitar párrafos vacíos del cuerpo
body = doc.element.body
for child in list(body):
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        text = ''.join(r.text or '' for r in child.findall('.//' + qn('w:t')))
        if not text.strip():
            body.remove(child)

# ─────────────────────────────────────────────
# PORTADA
# ─────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('INFORME DE SEGUIMIENTO\nCENTRO DE SERVICIOS — P0321')
run.font.name = 'Montserrat'
run.font.size = Pt(20)
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 0)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Bogotá, 14 de abril de 2026')
run2.font.name = 'Montserrat'
run2.font.size = Pt(12)
run2.font.color.rgb = RGBColor(0, 0, 0)

doc.add_page_break()

# ─────────────────────────────────────────────
# 1. INTRODUCCIÓN
# ─────────────────────────────────────────────
set_titulo(doc, '1. Introducción')
set_parrafo(doc,
    'El presente informe consolida el seguimiento al tablero de tareas del equipo de Centro de Servicios '
    '(Proyecto P0321 – Soporte Interno) de Blend 360, Engineering Colombia. La información fue obtenida '
    'mediante el cruce entre la transcripción de la reunión diaria (Daily) celebrada el 10 de abril de 2026 '
    'y el tablero de actividades exportado de ClickUp con corte al 14 de abril de 2026.')
set_parrafo(doc,
    'El objetivo del informe es identificar el estado actual de las iniciativas, señalar las modificaciones '
    'de fechas, cambios de responsable y ajustes de prioridad acordados durante la reunión, y establecer '
    'un cronograma de trabajo claro para el equipo. Los responsables principales involucrados son: '
    'Johan Espino, Johan Velandia y Luis Hernando Bernal M., con apoyo de Jose Florez.')

# ─────────────────────────────────────────────
# 2. DESARROLLO
# ─────────────────────────────────────────────
set_titulo(doc, '2. Desarrollo')
set_parrafo(doc,
    'A continuación se presenta el análisis detallado de cada iniciativa activa, con sus respectivas '
    'observaciones de seguimiento derivadas del contraste entre el Daily del 10 de abril y el tablero ClickUp.')

# ── 2.1 Houndoc / Reportes Zammad ──
set_subtitulo(doc, '2.1 Houndoc para Reportes de Cliente (Tarea 2.4)')
tabla_datos = [
    ('Responsables', 'Johan Velandia, Johan Espino, Luis Hernando Bernal M.'),
    ('Fecha original', '20 de marzo de 2026'),
    ('Fecha actual', '14 de abril de 2026 (última modificación en reunión del 10/04)'),
    ('Prioridad', 'URGENTE'),
]
t = doc.add_table(rows=len(tabla_datos), cols=2)
t.style = 'Table Grid'
for i, (k, v) in enumerate(tabla_datos):
    cell_text(t.rows[i].cells[0], k, bold=True, size=10)
    cell_text(t.rows[i].cells[1], v, size=10)
    shade_cell(t.rows[i].cells[0], 'D9D9D9')
doc.add_paragraph()
set_parrafo(doc,
    'Esta iniciativa acumula múltiples postergaciones desde finales de febrero. En la reunión del 10 de abril '
    'se evidenció que no se han iniciado los esfuerzos. Las subtareas críticas son: creación de usuario de '
    'prueba y asignación de roles (14/04), migración de base de datos a Houndoc (16/04), construcción del '
    'dashboard (22/04), elaboración de manuales (24/04), notificación formal (27/04) y capacitación a clientes '
    '(17/04). La gerencia destacó la urgencia de iniciar de inmediato dado que el retraso compromete la '
    'entrega de valor al cliente.')

# ── 2.2 Data Center ──
set_subtitulo(doc, '2.2 Data Center — Organización y Cableado (Tarea 3.2)')
tabla_datos2 = [
    ('Responsables', 'Johan Velandia, Johan Espino, Luis Hernando Bernal M.'),
    ('Fecha original', 'Múltiples postergaciones desde inicio de año'),
    ('Fecha actual', '17–21 de abril de 2026'),
    ('Prioridad', 'URGENTE'),
    ('Observación', 'Se mencionó la posibilidad de horas extra remuneradas un sábado'),
]
t2 = doc.add_table(rows=len(tabla_datos2), cols=2)
t2.style = 'Table Grid'
for i, (k, v) in enumerate(tabla_datos2):
    cell_text(t2.rows[i].cells[0], k, bold=True, size=10)
    cell_text(t2.rows[i].cells[1], v, size=10)
    shade_cell(t2.rows[i].cells[0], 'D9D9D9')
doc.add_paragraph()
set_parrafo(doc,
    'El data center presenta desorden persistente con cables sin marquillar y equipos sin clasificar. '
    'Existen bandejas adquiridas que no han sido instaladas (3.6). En la reunión se acordó: (i) organizar '
    'el espacio físico del almacén, (ii) realizar inventario y depuración de equipos (separar funcionales '
    'de no funcionales), (iii) instalar bandejas y (iv) marquillar cables. Se evaluó la opción de horas '
    'extra remuneradas un sábado para completar estas actividades. La tarea de inventario y almacenamiento '
    '(3.5) tiene fecha límite 20/04.')

# ── 2.3 Depuración Correos NUVU ──
set_subtitulo(doc, '2.3 Depuración de Correos con Dominio nuvu.cc (Tareas 8.2 y 8.4)')
tabla_datos3 = [
    ('Responsables', 'Johan Velandia, Luis Hernando Bernal M.'),
    ('Fecha original', '27 de febrero de 2026'),
    ('Fecha actual', '15 de abril de 2026 (miércoles)'),
    ('Prioridad', 'URGENTE'),
    ('Observación', 'Cuentas con tarjeta Clara no serán eliminadas; se requiere validación con Edwin y Nicolás Revollo'),
]
t3 = doc.add_table(rows=len(tabla_datos3), cols=2)
t3.style = 'Table Grid'
for i, (k, v) in enumerate(tabla_datos3):
    cell_text(t3.rows[i].cells[0], k, bold=True, size=10)
    cell_text(t3.rows[i].cells[1], v, size=10)
    shade_cell(t3.rows[i].cells[0], 'D9D9D9')
doc.add_paragraph()
set_parrafo(doc,
    'La tarea acumula casi dos meses de retraso. Johan Velandia recibió el listado de cuentas a depurar y '
    'dejó un comentario sobre cuentas con tarjeta Clara, pero no ejecutó la acción. En la reunión del 10/04 '
    'se determinó: (i) Johan Velandia debe contactar a Edwin y Nicolás Revollo (por mensaje o correo) para '
    'validar sus cuentas antes del 15/04, (ii) las demás cuentas sin observación se eliminarán directamente. '
    'Luis Hernando Bernal M. fue asignado como responsable adicional para garantizar la alerta de vencimiento. '
    'También se incluye depuración de Drive (8.4) con la misma fecha límite.')

# ── 2.4 WhatsApp Business ──
set_subtitulo(doc, '2.4 WhatsApp Business — Desarrollo N8N (Tarea 5.2)')
tabla_datos4 = [
    ('Responsable', 'Johan Espino'),
    ('Fecha original', 'Principios de febrero 2026'),
    ('Fecha actual', '30 de abril de 2026'),
    ('Prioridad', 'ALTA'),
    ('Bloqueo', 'Cuenta de META sin permisos correctos; requiere gestión con Jose Florez'),
]
t4 = doc.add_table(rows=len(tabla_datos4), cols=2)
t4.style = 'Table Grid'
for i, (k, v) in enumerate(tabla_datos4):
    cell_text(t4.rows[i].cells[0], k, bold=True, size=10)
    cell_text(t4.rows[i].cells[1], v, size=10)
    shade_cell(t4.rows[i].cells[0], 'D9D9D9')
doc.add_paragraph()
set_parrafo(doc,
    'La tarea de solicitar acceso a la plantilla de META (5.1) tiene fecha 13/04 y está bloqueada porque '
    'Jose Florez asignó permisos a una cuenta inexistente. En la reunión se acordó que Johan Espino debe '
    'enviar un correo formal a Jose Florez para que corrija el acceso, ya que el correo garantiza trazabilidad '
    'del seguimiento. El desarrollo N8N (5.2) permanece pendiente hasta resolver el bloqueo de acceso.')

# ── 2.5 Alarmas Cloud en Zammad ──
set_subtitulo(doc, '2.5 Alarmas de Cloud en Zammad (Tarea 13)')
tabla_datos5 = [
    ('Responsables', 'Johan Espino, Luis Hernando Bernal M.'),
    ('Subtareas críticas', 'Pruebas de carga (13.3): jueves 16/04 | Envío alarmas test (13.4): viernes 17/04 | Réplica producción (13.5): 21/04'),
    ('Prioridad', 'URGENTE'),
    ('Condición', 'Requiere estabilizar Zammad previamente; actualización Kubernetes pendiente'),
]
t5 = doc.add_table(rows=len(tabla_datos5), cols=2)
t5.style = 'Table Grid'
for i, (k, v) in enumerate(tabla_datos5):
    cell_text(t5.rows[i].cells[0], k, bold=True, size=10)
    cell_text(t5.rows[i].cells[1], v, size=10)
    shade_cell(t5.rows[i].cells[0], 'D9D9D9')
doc.add_paragraph()
set_parrafo(doc,
    'En la reunión del 10/04 se estableció una secuencia de ejecución: primero optimizar la estabilidad de '
    'Zammad (que presenta inestabilidad por aumento de base de conocimientos), luego realizar las pruebas de '
    'carga el jueves 16/04, envío de alarmas a ambiente de prueba el viernes 17/04, y réplica en producción '
    'la semana del 21/04. La configuración N8N para recepción de alarmas (13.6.1) tiene fecha límite 15/04. '
    'Johan Espino indicó que los workflows están en desarrollo. Se requiere también actualizar la versión de '
    'Kubernetes antes de que expire su soporte.')

# ── 2.6 Suite Microsoft ──
set_subtitulo(doc, '2.6 Suite Microsoft — Depuración Drive nuvu.cc (Tarea 8.4) y Correos (8.2)')
set_parrafo(doc,
    'Ambas tareas tienen fecha límite 15/04 con prioridad URGENTE. La depuración de correos y drive del '
    'dominio nuvu.cc está bloqueada por la misma razón que la tarea 8.2. Se coordina en conjunto con la '
    'depuración de cuentas de correo.')

# ── 2.7 QuickSight ──
set_subtitulo(doc, '2.7 Revisión y Depuración Tablero QuickSight (Tarea 15)')
tabla_datos7 = [
    ('Responsable principal', 'Luis Hernando Bernal M.'),
    ('Fecha', '30 de abril de 2026'),
    ('Prioridad', 'URGENTE'),
    ('Alcance', '150 tableros en cuenta 10; coordinar depuración con cada PM de proyecto'),
]
t7 = doc.add_table(rows=len(tabla_datos7), cols=2)
t7.style = 'Table Grid'
for i, (k, v) in enumerate(tabla_datos7):
    cell_text(t7.rows[i].cells[0], k, bold=True, size=10)
    cell_text(t7.rows[i].cells[1], v, size=10)
    shade_cell(t7.rows[i].cells[0], 'D9D9D9')
doc.add_paragraph()
set_parrafo(doc,
    'Se identificaron 150 tableros en la cuenta 10 de QuickSight, muchos en desuso o pertenecientes a '
    'proyectos de clientes que deberían estar en sus propias cuentas (Policía, Bioagraria, entre otros). '
    'En la reunión se acordó que Luis Hernando Bernal M. realizará el inventario y revisión de EC2 (15.1), '
    'y luego coordinará con cada PM la decisión de migrar o eliminar los tableros (15.2). '
    'El proceso es extenso y requiere apoyo del equipo de Luisa para las migraciones.')

# ── 2.8 Otros ──
set_subtitulo(doc, '2.8 Otras Iniciativas')
set_parrafo(doc, 'Diseño Proceso Seguridad Informática (Tarea 10): Fecha 13/04, prioridad URGENTE. '
    'Luis Hernando Bernal M. debe entregar un paquete de avance antes de la próxima sesión con el sponsor.')
set_parrafo(doc, 'Nuevos Ingresos Lunes 13/04 (Tarea 16): Johan Espino y Johan Velandia deben preparar '
    'equipos; se mencionó el envío de un equipo a Barranquilla.')
set_parrafo(doc, 'Empalme (Tarea 14): Transferencia de conocimiento entre Johan Espino y Johan Velandia '
    'para garantizar backup mutuo. Fecha 14/04, prioridad ALTA.')
set_parrafo(doc, 'Cierre ServiceCenter (Tarea 1.5.1.5): Validar viabilidad de migrar casos al nuevo sistema '
    'GitHub. Johan Espino debe contactar a Jose Florez antes del viernes 17/04 para confirmar si el '
    'proceso es viable.')

# ─────────────────────────────────────────────
# 3. COMPROMISOS
# ─────────────────────────────────────────────
set_titulo(doc, '3. Compromisos')
set_parrafo(doc, 'Los siguientes compromisos fueron acordados explícitamente durante la reunión del 10 de abril de 2026:')

compromisos = [
    ('Johan Espino',         'Enviar correo formal a Jose Florez solicitando corrección de acceso META para WhatsApp Business.',    '10/04/2026'),
    ('Johan Velandia',       'Contactar a Edwin y Nicolás Revollo para validar cuentas antes de depuración nuvu.cc.',               '13/04/2026'),
    ('Johan Espino / Johan Velandia',  'Preparar equipos para nuevos ingresos del lunes 13 de abril; enviar equipo a Barranquilla.',       '13/04/2026'),
    ('Johan Espino',                   'Completar configuración N8N para recepción de alarmas (13.6.1).',                             '15/04/2026'),
    ('Johan Velandia / Luis H. Bernal M.', 'Completar depuración correos y drive nuvu.cc (8.2 y 8.4).',                              '15/04/2026'),
    ('Johan Espino',                   'Ejecutar pruebas de carga y estrés en Zammad (13.3).',                                        '16/04/2026'),
    ('Johan Espino',                   'Envío de alarmas a ambiente de prueba (13.4).',                                               '17/04/2026'),
    ('Johan Espino',                   'Validar viabilidad migración casos ServiceCenter con Jose Florez (1.5.1.5).',                 '17/04/2026'),
    ('Johan Espino / Johan Velandia',  'Organizar espacio Data Center: depuración, inventario, instalación de bandejas (3.4, 3.5, 3.6).', '17/04–21/04'),
    ('Johan Espino / Johan Velandia',  'Crear usuario de prueba en Houndoc y asignar roles (2.3.1).',                                '14/04/2026'),
    ('Luis H. Bernal M.',              'Iniciar revisión de 150 tableros QuickSight cuenta 10 (15.1).',                               '30/04/2026'),
    ('Luis H. Bernal M.',              'Entregar avance de diseño proceso seguridad informática (Tarea 10).',                          '13/04/2026'),
    ('Johan Espino',         'Realizar empalme de conocimiento con Johan Velandia (Tarea 14).',                                    '14/04/2026'),
]

t_comp = doc.add_table(rows=1 + len(compromisos), cols=3)
t_comp.style = 'Table Grid'
headers_comp = ['Responsable', 'Compromiso', 'Fecha Límite']
for j, h in enumerate(headers_comp):
    cell_text(t_comp.rows[0].cells[j], h, bold=True, size=10, color='FFFFFF')
    shade_cell(t_comp.rows[0].cells[j], '1F3864')
for i, (resp, comp, fecha) in enumerate(compromisos):
    cell_text(t_comp.rows[i+1].cells[0], resp, size=10)
    cell_text(t_comp.rows[i+1].cells[1], comp, size=10)
    cell_text(t_comp.rows[i+1].cells[2], fecha, size=10)
    if i % 2 == 0:
        for j in range(3):
            shade_cell(t_comp.rows[i+1].cells[j], 'EBF0FF')

# ─────────────────────────────────────────────
# 4. PENDIENTES
# ─────────────────────────────────────────────
set_titulo(doc, '4. Pendientes')
set_parrafo(doc, 'Se identifican las siguientes tareas con mayor riesgo de incumplimiento o que requieren atención inmediata:')

pendientes = [
    ('Houndoc para reportes de cliente', '2.4', 'URGENTE', 'Johan Espino / Johan Velandia', 'Múltiples postergaciones desde feb. Sin inicio reportado.'),
    ('Organización Data Center', '3.2', 'URGENTE', 'Johan Espino / Johan Velandia / Luis H. Bernal M.', 'Pendiente desde inicio de año. Bandejas sin instalar.'),
    ('Depuración correos nuvu.cc', '8.2 / 8.4', 'URGENTE', 'Johan Velandia / Luis H. Bernal M.', 'Vencida desde 27 feb. Requiere validación con 2 personas.'),
    ('WhatsApp Business N8N', '5.2', 'ALTA', 'Johan Espino', 'Bloqueado por acceso META no otorgado correctamente.'),
    ('Zammad estabilidad + alarmas', '13.3 – 13.5', 'URGENTE', 'Johan Espino / Luis H. Bernal M.', 'Dependencia: debe estabilizarse Zammad antes de proceder.'),
    ('Diseño proceso seguridad informática', '10', 'URGENTE', 'Luis H. Bernal M.', 'Sin avances documentados para presentar al sponsor.'),
    ('Revisión tableros QuickSight', '15.1 / 15.2', 'URGENTE', 'Luis H. Bernal M.', '150 tableros por revisar. Tarea de largo aliento.'),
    ('Publicación casos ServiceCenter', '1.5.1.5', 'NORMAL', 'Johan Espino / Jose Florez', 'Definir viabilidad con Jose Florez antes del 17/04.'),
    ('Fortinet — 4 subtareas', '4.1 – 4.4', 'NORMAL/ALTA', 'Johan Velandia / Luis H. Bernal M.', 'Sin fecha definida en varias subtareas. Sin mención en daily.'),
    ('Babybot — Arquitectura y manuales', '13.2', 'NORMAL', 'Johan Espino', 'Vencida desde feb. Sin mención en daily.'),
]

t_pend = doc.add_table(rows=1 + len(pendientes), cols=5)
t_pend.style = 'Table Grid'
headers_p = ['Tarea', 'ID', 'Prioridad', 'Responsable', 'Observación']
for j, h in enumerate(headers_p):
    cell_text(t_pend.rows[0].cells[j], h, bold=True, size=9, color='FFFFFF')
    shade_cell(t_pend.rows[0].cells[j], '1F3864')
for i, row in enumerate(pendientes):
    for j, val in enumerate(row):
        cell_text(t_pend.rows[i+1].cells[j], val, size=9)
    prioridad = row[2]
    color_fila = PRIORIDAD_COLOR.get(prioridad.split('/')[0].strip(), 'FFFFFF')
    for j in range(5):
        shade_cell(t_pend.rows[i+1].cells[j], color_fila)

# ─────────────────────────────────────────────
# 5. COMENTARIOS
# ─────────────────────────────────────────────
set_titulo(doc, '5. Comentarios')

set_parrafo(doc,
    'La reunión del 10 de abril reveló un patrón recurrente de postergación de tareas, especialmente en '
    'iniciativas de alta prioridad que llevan más de dos meses sin ejecutarse. Se destacan las siguientes '
    'observaciones de gestión:')

comentarios = [
    'Trazabilidad insuficiente: Múltiples actividades se realizan sin dejar evidencia documental (capturas, informes, comentarios en ClickUp). La gerencia enfatizó la necesidad de registrar cada acción como soporte de gestión.',
    'Priorización reactiva: El equipo atiende urgencias del día a día (onboardings, equipos, redes) sin avanzar en las iniciativas estratégicas. Se requiere una metodología de gestión del tiempo que proteja los bloques de trabajo planificado.',
    'Dependencias externas no gestionadas: Tareas como WhatsApp Business (META) y casos ServiceCenter dependen de terceros (Jose Florez) y no se han escalado formalmente a través de correo, lo que reduce la trazabilidad del bloqueo.',
    'Riesgo de concentración de conocimiento: Johan Espino concentra la mayoría de las tareas críticas. El empalme con Johan Velandia es urgente para garantizar continuidad operativa.',
    'Data Center como deuda técnica creciente: La postergación sistemática de la organización del data center aumenta el riesgo operativo. Se recomienda programar una jornada específica (sábado con reconocimiento económico) para saldar esta deuda.',
    'Zammad bajo presión: El crecimiento de la base de conocimiento está generando inestabilidad. La actualización de Kubernetes es inminente y no debe postergarse más allá de las pruebas de alarmas.',
]
for c in comentarios:
    set_bullet(doc, c)

# ─────────────────────────────────────────────
# 6. CRONOGRAMA DE TRABAJO — ABRIL 2026
# ─────────────────────────────────────────────
from docx.shared import Cm
from docx.oxml import OxmlElement

set_titulo(doc, '6. Cronograma de Trabajo — Abril 2026')
set_parrafo(doc,
    'Cronograma consolidado de tareas activas — Proyecto P0321 Soporte Interno. '
    'Ordenado de fecha más reciente a más antigua, agrupado por semana de trabajo. '
    'Fuente: ClickUp P0321, corte 14 de abril de 2026.')

# ── Leyenda de prioridades ──
doc.add_paragraph()
p_leyenda = doc.add_paragraph()
p_leyenda.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_leg = p_leyenda.add_run('LEYENDA DE PRIORIDADES:  ')
run_leg.font.name = 'Montserrat'; run_leg.font.size = Pt(9); run_leg.bold = True

for label, color_hex in [('URGENTE', COLOR_URGENTE), ('ALTA', COLOR_HIGH),
                          ('NORMAL', COLOR_NORMAL), ('BAJA', COLOR_LOW)]:
    t_leg = doc.add_table(rows=1, cols=2)
    t_leg.style = 'Table Grid'
    t_leg.rows[0].cells[0].width = Cm(0.6)
    t_leg.rows[0].cells[1].width = Cm(2.0)
    shade_cell(t_leg.rows[0].cells[0], color_hex.lstrip('#') if color_hex.startswith('#') else color_hex)
    cell_text(t_leg.rows[0].cells[1], label, size=8, bold=True)

# ── Tabla leyenda compacta en una sola fila ──
doc.add_paragraph()
t_ley = doc.add_table(rows=2, cols=8)
t_ley.style = 'Table Grid'
leyenda_items = [
    ('', 'URGENTE', COLOR_URGENTE),
    ('', 'ALTA',    COLOR_HIGH),
    ('', 'NORMAL',  COLOR_NORMAL),
    ('', 'BAJA',    COLOR_LOW),
]
for idx, (_, lbl, col) in enumerate(leyenda_items):
    shade_cell(t_ley.rows[0].cells[idx*2],     col)
    cell_text( t_ley.rows[0].cells[idx*2],     '  ', size=8)
    cell_text( t_ley.rows[0].cells[idx*2 + 1], lbl,  size=8, bold=True)
    shade_cell(t_ley.rows[1].cells[idx*2],     col)
    cell_text( t_ley.rows[1].cells[idx*2],     '  ', size=8)
    cell_text( t_ley.rows[1].cells[idx*2 + 1], '',   size=8)

# Reemplazar leyenda por una limpia en una sola tabla de 1 fila x 8 celdas
for tbl in [t_ley]:
    tbl._element.getparent().remove(tbl._element)

doc.add_paragraph()
t_leyenda = doc.add_table(rows=1, cols=8)
t_leyenda.style = 'Table Grid'
leyenda_cols = [
    ('URGENTE', 'FFCCCC'), ('ALTA', 'FFE5CC'),
    ('NORMAL',  'CCE0FF'), ('BAJA', 'D9F0CC'),
]
for idx, (lbl, col) in enumerate(leyenda_cols):
    shade_cell(t_leyenda.rows[0].cells[idx*2],     col)
    cell_text( t_leyenda.rows[0].cells[idx*2],     '  ', size=9)
    cell_text( t_leyenda.rows[0].cells[idx*2 + 1], lbl,  size=9, bold=True)

doc.add_paragraph()

# ── Datos del cronograma (descendente: más reciente primero) ──
# Columnas: #, Fecha, ID, Tarea, Responsable, Prioridad, Iniciativa
cronograma = [
    # SEMANA 28 ABR – 30 ABR
    ('30/04', '15.2',   'Coordinar depuración tableros QuickSight con PMs',          'Luis H. Bernal M.',                              'NORMAL',  'QuickSight'),
    ('30/04', '15.1',   'Revisión EC2 tablero consola Blend QuickSight',             'Luis H. Bernal M. / Johan Espino',               'URGENTE', 'QuickSight'),
    ('30/04', '5.2',    'Desarrollar N8N para WhatsApp Business',                    'Johan Espino',                                   'ALTA',    'WhatsApp'),
    ('27/04', '2.3.5',  'Notificación formal Houndoc a clientes',                    'Luis H. Bernal M.',                              'URGENTE', 'Houndoc'),
    # SEMANA 21 ABR – 25 ABR
    ('24/04', '2.3.4',  'Construir manuales Houndoc',                                'Johan Espino / Johan Velandia',                  'URGENTE', 'Houndoc'),
    ('22/04', '2.3.3',  'Construir Dashboard Houndoc y realizar pruebas',            'Johan Espino / Johan Velandia',                  'URGENTE', 'Houndoc'),
    ('21/04', '13.7',   'Sesión definición flujo de alarmas cloud y CS',             'Luis H. Bernal M.',                              'URGENTE', 'Zammad / Alarmas'),
    ('21/04', '13.5',   'Réplica de alarmas en producción',                          'Johan Espino',                                   'URGENTE', 'Zammad / Alarmas'),
    ('21/04', '3.2',    'Organización cables y marquillado Data Center',             'Johan Espino / Johan Velandia',                  'URGENTE', 'Data Center'),
    ('20/04', '3.5',    'Inventario y almacenamiento de equipos',                    'Luis H. Bernal M. / Johan Velandia / Johan Espino', 'URGENTE', 'Data Center'),
    # SEMANA 14 ABR – 17 ABR
    ('17/04', '2.3.6',  'Capacitación a clientes — Houndoc',                        'Luis H. Bernal M. / Johan Velandia / Johan Espino', 'URGENTE', 'Houndoc'),
    ('17/04', '1.5.1.5','Validar viabilidad migración casos ServiceCenter',          'Johan Espino / Jose Florez',                     'NORMAL',  'ServiceCenter'),
    ('17/04', '13.4',   'Envío de alarmas a ambiente de prueba',                     'Johan Espino',                                   'URGENTE', 'Zammad / Alarmas'),
    ('17/04', '3.6',    'Instalación de bandejas Data Center',                       'Luis H. Bernal M. / Johan Velandia / Johan Espino', 'URGENTE', 'Data Center'),
    ('17/04', '3.4',    'Depuración de equipos Data Center',                         'Luis H. Bernal M. / Johan Velandia / Johan Espino', 'URGENTE', 'Data Center'),
    ('16/04', '13.3',   'Pruebas de carga y estrés con agente de alarmas',           'Johan Espino',                                   'URGENTE', 'Zammad / Alarmas'),
    ('16/04', '2.3.2',  'Migrar base de datos de usuarios a Houndoc',                'Johan Espino / Johan Velandia',                  'URGENTE', 'Houndoc'),
    ('15/04', '13.6.1', 'Configuración N8N para recepción de alarmas',               'Johan Espino',                                   'URGENTE', 'Zammad / Alarmas'),
    ('15/04', '12.6',   'Revisar estabilidad Zammad',                                'Johan Espino / Luis H. Bernal M.',               'URGENTE', 'Zammad / Alarmas'),
    ('15/04', '8.4',    'Depuración Drive dominio nuvu.cc',                          'Johan Velandia / Luis H. Bernal M.',             'URGENTE', 'Suite Microsoft'),
    ('15/04', '8.2',    'Depuración correos dominio nuvu.cc',                        'Johan Velandia / Luis H. Bernal M.',             'URGENTE', 'Suite Microsoft'),
    ('14/04', '14',     'Empalme de conocimiento Johan Espino → Johan Velandia',     'Johan Espino / Johan Velandia',                  'ALTA',    'Empalme'),
    ('14/04', '2.3.1',  'Crear usuario de prueba y asignar roles (Houndoc)',         'Johan Espino / Johan Velandia',                  'URGENTE', 'Houndoc'),
]

# Semanas (descendente)
semanas = [
    ('SEMANA 28 – 30 ABR 2026',  ['30/04', '27/04']),
    ('SEMANA 21 – 25 ABR 2026',  ['21/04', '22/04', '24/04', '20/04']),
    ('SEMANA 14 – 17 ABR 2026',  ['14/04', '15/04', '16/04', '17/04']),
]

COLS = ['#', 'Fecha', 'ID', 'Tarea / Actividad', 'Responsable', 'Prioridad', 'Iniciativa']
COL_WIDTHS = [Cm(0.7), Cm(1.4), Cm(1.1), Cm(6.5), Cm(3.5), Cm(1.8), Cm(2.8)]

def set_col_widths(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

contador = 1
for sem_label, fechas_sem in semanas:
    tareas_sem = [t for t in cronograma if t[0] in fechas_sem]

    # Header de semana
    t_sem = doc.add_table(rows=1, cols=1)
    t_sem.style = 'Table Grid'
    cell_text(t_sem.rows[0].cells[0], f'  {sem_label}', bold=True, size=11, color='FFFFFF')
    shade_cell(t_sem.rows[0].cells[0], '1F3864')

    # Header de columnas
    t_head = doc.add_table(rows=1, cols=len(COLS))
    t_head.style = 'Table Grid'
    set_col_widths(t_head, COL_WIDTHS)
    for j, h in enumerate(COLS):
        cell_text(t_head.rows[0].cells[j], h, bold=True, size=9, color='FFFFFF')
        shade_cell(t_head.rows[0].cells[j], '2E4057')

    # Filas de datos
    t_data = doc.add_table(rows=len(tareas_sem), cols=len(COLS))
    t_data.style = 'Table Grid'
    set_col_widths(t_data, COL_WIDTHS)

    fecha_ant = None
    for i, (fecha, tid, tarea, resp, prio, inic) in enumerate(tareas_sem):
        color_fila = PRIORIDAD_COLOR.get(prio, 'FFFFFF')
        # Indicador visual de cambio de fecha
        fecha_display = fecha if fecha != fecha_ant else ''
        fecha_ant = fecha

        cell_text(t_data.rows[i].cells[0], str(contador),    size=9, bold=True)
        cell_text(t_data.rows[i].cells[1], fecha_display,    size=9, bold=(fecha_display != ''))
        cell_text(t_data.rows[i].cells[2], tid,              size=9)
        cell_text(t_data.rows[i].cells[3], tarea,            size=9)
        cell_text(t_data.rows[i].cells[4], resp,             size=9)
        cell_text(t_data.rows[i].cells[5], prio,             size=9, bold=True, color='FFFFFF')
        cell_text(t_data.rows[i].cells[6], inic,             size=9)

        for j in range(len(COLS)):
            shade_cell(t_data.rows[i].cells[j], color_fila)
        # Celda de fecha con color más intenso si cambia
        if fecha_display:
            shade_cell(t_data.rows[i].cells[1], 'D0D8E8')
        # Mapa de calor: celda Prioridad con color sólido
        shade_cell(t_data.rows[i].cells[5], PRIORIDAD_HEAT.get(prio, '888888'))

        contador += 1

    doc.add_paragraph()

# ── Resumen por responsable ──
set_subtitulo(doc, 'Resumen de Carga por Responsable')
set_parrafo(doc, 'Distribución de tareas del cronograma de abril 2026 por responsable:')

from collections import Counter
resp_conteo = Counter()
for _, _, tarea, resp, prio, inic in cronograma:
    for r in resp.split('/'):
        resp_conteo[r.strip()] += 1

t_resp = doc.add_table(rows=1 + len(resp_conteo), cols=3)
t_resp.style = 'Table Grid'
headers_r = ['Responsable', 'N.° de Tareas', 'Observación']
for j, h in enumerate(headers_r):
    cell_text(t_resp.rows[0].cells[j], h, bold=True, size=10, color='FFFFFF')
    shade_cell(t_resp.rows[0].cells[j], '1F3864')

obs_resp = {
    'Johan Espino':      'Responsable con mayor carga. Requiere apoyo urgente.',
    'Johan Velandia':    'Backup de Johan Espino. Data Center + Suite Microsoft.',
    'Luis H. Bernal M.': 'QuickSight, Zammad, Notificaciones, Seguridad.',
    'Jose Florez':       'Soporte en ServiceCenter y WhatsApp META.',
}
for i, (resp_n, cnt) in enumerate(sorted(resp_conteo.items(), key=lambda x: -x[1])):
    cell_text(t_resp.rows[i+1].cells[0], resp_n, size=10)
    cell_text(t_resp.rows[i+1].cells[1], str(cnt), size=10, bold=True)
    cell_text(t_resp.rows[i+1].cells[2], obs_resp.get(resp_n, ''), size=10)
    if i % 2 == 0:
        for j in range(3):
            shade_cell(t_resp.rows[i+1].cells[j], 'EBF0FF')

# Guardar
doc.save(SALIDA)
print(f'Documento generado: {SALIDA}')
print(f'Total tareas en cronograma: {len(cronograma)}')
