"""
Generador de Resumen Ejecutivo - Centro de Servicios
Salida: Resumen_Ejecutivo_CentroServicios_20260414.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

BASE     = Path(r'c:\Users\BOG-LAP-SER-176\Documents\bld-engineering-pmo-col\Centro de Servicios')
OUTPUT   = BASE / '3. Output'
PLANTILLA = BASE / '1. Input' / 'Plantilla word Blend.docx'
SALIDA   = OUTPUT / 'Resumen_Ejecutivo_CentroServicios_20260414.docx'
OUTPUT.mkdir(exist_ok=True)

# ── Colores corporativos ──
AZUL_DARK  = '1F3864'
AZUL_MED   = '2E4057'
AZUL_LIGHT = 'D6E4F7'
GRIS       = 'F2F2F2'
ROJO       = 'C00000'
VERDE      = '375623'
NARANJA    = 'BF5900'

PRIORIDAD_COLOR = {
    'URGENTE': 'FFCCCC',
    'ALTA':    'FFE5CC',
    'NORMAL':  'CCE0FF',
    'BAJA':    'D9F0CC',
}
PRIORIDAD_TEXTO = {
    'URGENTE': 'C00000',
    'ALTA':    'BF5900',
    'NORMAL':  '1F3864',
    'BAJA':    '375623',
}

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_text(cell, texto, bold=False, size=10, color='000000', align='LEFT'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = getattr(WD_ALIGN_PARAGRAPH, align)
    run = p.add_run(texto)
    run.font.name = 'Montserrat'
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(
        int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))

def titulo(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(texto)
    run.font.name = 'Montserrat'
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    return p

def parrafo(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(texto)
    run.font.name = 'Montserrat'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def bullet(doc, texto, bold_label=None):
    p = doc.add_paragraph(style='List Paragraph')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_label:
        r1 = p.add_run(f'• {bold_label}: ')
        r1.font.name = 'Montserrat'; r1.font.size = Pt(10); r1.bold = True
        r1.font.color.rgb = RGBColor(0, 0, 0)
        r2 = p.add_run(texto)
        r2.font.name = 'Montserrat'; r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run = p.add_run(f'• {texto}')
        run.font.name = 'Montserrat'; run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

# ── Crear documento desde plantilla ──
doc = Document(PLANTILLA)
body = doc.element.body
for child in list(body):
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        text = ''.join(r.text or '' for r in child.findall('.//' + qn('w:t')))
        if not text.strip():
            body.remove(child)
for p in doc.paragraphs:
    p.clear()
for t in doc.tables:
    t._element.getparent().remove(t._element)

# ═══════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════
p_port = doc.add_paragraph()
p_port.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_port.add_run('RESUMEN EJECUTIVO\nCENTRO DE SERVICIOS — P0321\nSoporte Interno')
r.font.name = 'Montserrat'; r.font.size = Pt(18); r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

p_fecha = doc.add_paragraph()
p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_fecha.add_run('Bogotá, 14 de abril de 2026  |  Blend 360 – Engineering Colombia')
r2.font.name = 'Montserrat'; r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 1. INDICADORES CLAVE
# ═══════════════════════════════════════════════════
titulo(doc, '1. Indicadores Clave del Periodo')
doc.add_paragraph()

kpis = [
    ('23',         'Tareas activas\nabriles',       AZUL_DARK),
    ('15',         'Tareas\nURGENTE',               'C00000'),
    ('3',          'Tareas\nBLOQUEADAS',            '7030A0'),
    ('3',          'Responsables\nprincipales',      '375623'),
    ('2+ meses',   'Retraso máximo\nacumulado',     'BF5900'),
    ('6',          'Iniciativas\nactivas',           '1F3864'),
]

t_kpi = doc.add_table(rows=2, cols=len(kpis))
t_kpi.style = 'Table Grid'
for j, (valor, label, color) in enumerate(kpis):
    # Fila valor
    c_val = t_kpi.rows[0].cells[j]
    c_val.text = ''
    p_v = c_val.paragraphs[0]
    p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_v = p_v.add_run(valor)
    r_v.font.name = 'Montserrat'; r_v.font.size = Pt(22); r_v.bold = True
    r_v.font.color.rgb = RGBColor(
        int(color[0:2],16), int(color[2:4],16), int(color[4:6],16))
    shade_cell(c_val, GRIS)
    # Fila label
    c_lbl = t_kpi.rows[1].cells[j]
    c_lbl.text = ''
    p_l = c_lbl.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_l = p_l.add_run(label)
    r_l.font.name = 'Montserrat'; r_l.font.size = Pt(8); r_l.bold = False
    r_l.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    shade_cell(c_lbl, 'FFFFFF')

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 2. ESTADO POR INICIATIVA
# ═══════════════════════════════════════════════════
titulo(doc, '2. Estado por Iniciativa')
doc.add_paragraph()

iniciativas = [
    # (Iniciativa, Estado, Prioridad, Resp., Próx. entrega, Observación)
    ('Houndoc / Reportes', '🔴 EN RIESGO', 'URGENTE',
     'Johan Espino / Johan Velandia / Luis H. Bernal M.',
     '14 – 27 abr',
     'Postergado desde feb. Sin inicio a la fecha del Daily.'),
    ('Data Center', '🔴 EN RIESGO', 'URGENTE',
     'Johan Espino / Johan Velandia / Luis H. Bernal M.',
     '17 – 21 abr',
     'Deuda técnica de inicio de año. Bandejas sin instalar.'),
    ('Zammad + Alarmas', '🟡 EN PROGRESO', 'URGENTE',
     'Johan Espino / Luis H. Bernal M.',
     '15 – 21 abr',
     'Condicionado a estabilizar Zammad. Kubernetes pendiente.'),
    ('Suite Microsoft / nuvu.cc', '🔴 EN RIESGO', 'URGENTE',
     'Johan Velandia / Luis H. Bernal M.',
     '15 abr',
     'Vencida desde 27 feb. Requiere validación con 2 personas.'),
    ('WhatsApp Business', '⛔ BLOQUEADO', 'ALTA',
     'Johan Espino',
     '30 abr',
     'Acceso META sin configurar correctamente por Jose Florez.'),
    ('QuickSight', '🟡 EN PROGRESO', 'URGENTE',
     'Luis H. Bernal M.',
     '30 abr',
     '150 tableros por revisar y depurar con PMs.'),
    ('ServiceCenter / Migración', '🟡 EN PROGRESO', 'NORMAL',
     'Johan Espino / Jose Florez',
     '17 abr',
     'Validar viabilidad con Jose Florez antes del cierre de semana.'),
    ('Empalme de Conocimiento', '🟡 EN PROGRESO', 'ALTA',
     'Johan Espino → Johan Velandia',
     '14 abr',
     'Transferencia de conocimiento técnico entre ambos ingenieros.'),
    ('Seguridad Informática', '🔴 EN RIESGO', 'URGENTE',
     'Luis H. Bernal M.',
     '13 abr (vencida)',
     'Sin avances documentados. Sponsor a la espera.'),
]

cols_ini = ['Iniciativa', 'Estado', 'Prioridad', 'Responsable', 'Próx. Entrega', 'Observación']
t_ini = doc.add_table(rows=1 + len(iniciativas), cols=len(cols_ini))
t_ini.style = 'Table Grid'
col_w_ini = [Cm(3.2), Cm(2.4), Cm(1.8), Cm(3.2), Cm(2.2), Cm(5.0)]
for row in t_ini.rows:
    for i, cell in enumerate(row.cells):
        cell.width = col_w_ini[i]

for j, h in enumerate(cols_ini):
    cell_text(t_ini.rows[0].cells[j], h, bold=True, size=9, color='FFFFFF')
    shade_cell(t_ini.rows[0].cells[j], AZUL_DARK)

for i, (ini, estado, prio, resp, fecha, obs) in enumerate(iniciativas):
    cell_text(t_ini.rows[i+1].cells[0], ini,    size=9, bold=True)
    cell_text(t_ini.rows[i+1].cells[1], estado, size=9, bold=True,
              color=PRIORIDAD_TEXTO.get(prio, '000000'))
    cell_text(t_ini.rows[i+1].cells[2], prio,   size=9, bold=True,
              color=PRIORIDAD_TEXTO.get(prio, '000000'))
    cell_text(t_ini.rows[i+1].cells[3], resp,   size=9)
    cell_text(t_ini.rows[i+1].cells[4], fecha,  size=9, bold=True)
    cell_text(t_ini.rows[i+1].cells[5], obs,    size=9)
    col_bg = PRIORIDAD_COLOR.get(prio, 'FFFFFF')
    for j in range(len(cols_ini)):
        shade_cell(t_ini.rows[i+1].cells[j], col_bg if j != 1 else 'FFFFFF')

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 3. COMPROMISOS INMEDIATOS (próximos 5 días)
# ═══════════════════════════════════════════════════
titulo(doc, '3. Compromisos Inmediatos — 14 al 17 de Abril')
doc.add_paragraph()

compromisos_inm = [
    ('14 abr', 'Crear usuario de prueba y asignar roles Houndoc (2.3.1)',                    'Johan Espino / Johan Velandia',                    'URGENTE'),
    ('14 abr', 'Empalme de conocimiento técnico Johan Espino → Johan Velandia (14)',          'Johan Espino / Johan Velandia',                    'ALTA'),
    ('15 abr', 'Depuración correos y Drive nuvu.cc (8.2 y 8.4)',                              'Johan Velandia / Luis H. Bernal M.',               'URGENTE'),
    ('15 abr', 'Revisar estabilidad Zammad (12.6)',                                           'Johan Espino / Luis H. Bernal M.',                 'URGENTE'),
    ('15 abr', 'Configuración N8N recepción de alarmas (13.6.1)',                             'Johan Espino',                                     'URGENTE'),
    ('16 abr', 'Migrar base de datos usuarios a Houndoc (2.3.2)',                             'Johan Espino / Johan Velandia',                    'URGENTE'),
    ('16 abr', 'Pruebas de carga y estrés Zammad (13.3)',                                     'Johan Espino',                                     'URGENTE'),
    ('17 abr', 'Depuración equipos + instalación bandejas Data Center (3.4/3.6)',             'Luis H. Bernal M. / Johan Velandia / Johan Espino','URGENTE'),
    ('17 abr', 'Envío de alarmas a ambiente de prueba (13.4)',                                'Johan Espino',                                     'URGENTE'),
    ('17 abr', 'Validar viabilidad migración casos ServiceCenter (1.5.1.5)',                  'Johan Espino / Jose Florez',                       'NORMAL'),
    ('17 abr', 'Capacitación a clientes Houndoc (2.3.6)',                                     'Luis H. Bernal M. / Johan Velandia / Johan Espino','URGENTE'),
]

t_comp = doc.add_table(rows=1 + len(compromisos_inm), cols=4)
t_comp.style = 'Table Grid'
col_w_comp = [Cm(1.6), Cm(7.5), Cm(3.5), Cm(2.2)]
for row in t_comp.rows:
    for i, cell in enumerate(row.cells):
        cell.width = col_w_comp[i]

for j, h in enumerate(['Fecha', 'Compromiso', 'Responsable', 'Prioridad']):
    cell_text(t_comp.rows[0].cells[j], h, bold=True, size=9, color='FFFFFF')
    shade_cell(t_comp.rows[0].cells[j], AZUL_DARK)

fecha_ant = None
for i, (fecha, comp, resp, prio) in enumerate(compromisos_inm):
    col_bg = PRIORIDAD_COLOR.get(prio, 'FFFFFF')
    fecha_display = fecha if fecha != fecha_ant else ''
    fecha_ant = fecha
    cell_text(t_comp.rows[i+1].cells[0], fecha_display, size=9, bold=True,
              color=AZUL_DARK if fecha_display else '888888')
    cell_text(t_comp.rows[i+1].cells[1], comp,          size=9)
    cell_text(t_comp.rows[i+1].cells[2], resp,          size=9)
    cell_text(t_comp.rows[i+1].cells[3], prio,          size=9, bold=True,
              color=PRIORIDAD_TEXTO.get(prio, '000000'))
    for j in range(4):
        shade_cell(t_comp.rows[i+1].cells[j], col_bg)
    if fecha_display:
        shade_cell(t_comp.rows[i+1].cells[0], 'D0D8E8')

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 4. RIESGOS Y ALERTAS
# ═══════════════════════════════════════════════════
titulo(doc, '4. Riesgos y Alertas')
doc.add_paragraph()

riesgos = [
    ('CRÍTICO',  'Concentración de conocimiento',
     'Johan Espino es responsable único de la mayoría de tareas críticas. '
     'El empalme con Johan Velandia es urgente.'),
    ('CRÍTICO',  'Acumulación de deuda técnica — Data Center',
     'La postergación sistemática de la organización física del data center '
     'eleva el riesgo operativo. Cables desconectados pueden generar caídas de red.'),
    ('ALTO',     'Dependencia externa sin gestión formal — WhatsApp / META',
     'Johan Espino no ha enviado correo formal a Jose Florez para resolver el '
     'bloqueo de permisos. Sin trazabilidad del seguimiento.'),
    ('ALTO',     'Inestabilidad Zammad bajo crecimiento de carga',
     'El aumento de usuarios y base de conocimiento está generando inestabilidad. '
     'La actualización de Kubernetes es inminente y no debe postergarse.'),
    ('MEDIO',    'Tareas sin evidencia documental',
     'Múltiples actividades se ejecutan sin dejar capturas, informes o comentarios '
     'en ClickUp. Riesgo de no acreditar el trabajo ante la gerencia.'),
    ('MEDIO',    'Priorización reactiva del equipo',
     'El equipo atiende urgencias del día a día (onboardings, redes, equipos) '
     'sacrificando el avance de iniciativas estratégicas planificadas.'),
]

RIESGO_COLOR = {'CRÍTICO': 'FFCCCC', 'ALTO': 'FFE5CC', 'MEDIO': 'FFF2CC'}
RIESGO_TEXTO = {'CRÍTICO': 'C00000', 'ALTO': 'BF5900', 'MEDIO': '7F6000'}

t_riesgos = doc.add_table(rows=1 + len(riesgos), cols=3)
t_riesgos.style = 'Table Grid'
col_w_r = [Cm(2.0), Cm(4.5), Cm(11.3)]
for row in t_riesgos.rows:
    for i, cell in enumerate(row.cells):
        cell.width = col_w_r[i]

for j, h in enumerate(['Nivel', 'Riesgo', 'Descripción']):
    cell_text(t_riesgos.rows[0].cells[j], h, bold=True, size=9, color='FFFFFF')
    shade_cell(t_riesgos.rows[0].cells[j], AZUL_DARK)

for i, (nivel, nombre, desc) in enumerate(riesgos):
    cell_text(t_riesgos.rows[i+1].cells[0], nivel,  size=9, bold=True,
              color=RIESGO_TEXTO.get(nivel, '000000'))
    cell_text(t_riesgos.rows[i+1].cells[1], nombre, size=9, bold=True)
    cell_text(t_riesgos.rows[i+1].cells[2], desc,   size=9)
    for j in range(3):
        shade_cell(t_riesgos.rows[i+1].cells[j], RIESGO_COLOR.get(nivel, 'FFFFFF'))

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 5. CARGA POR RESPONSABLE
# ═══════════════════════════════════════════════════
titulo(doc, '5. Distribución de Carga por Responsable')
doc.add_paragraph()

carga = [
    ('Johan Espino',         12, 'URGENTE', 'Máxima carga. Riesgo de cuello de botella.'),
    ('Johan Velandia',        8, 'URGENTE', 'Backup de Espino. Data Center + Suite Microsoft.'),
    ('Luis H. Bernal M.',     6, 'URGENTE', 'QuickSight, Zammad, Notificaciones, Seguridad.'),
    ('Jose Florez',           2, 'NORMAL',  'Soporte en ServiceCenter y WhatsApp META.'),
]

t_carga = doc.add_table(rows=1 + len(carga), cols=4)
t_carga.style = 'Table Grid'
col_w_c = [Cm(4.0), Cm(2.5), Cm(2.5), Cm(8.8)]
for row in t_carga.rows:
    for i, cell in enumerate(row.cells):
        cell.width = col_w_c[i]

for j, h in enumerate(['Responsable', 'N.° Tareas', 'Nivel', 'Observación']):
    cell_text(t_carga.rows[0].cells[j], h, bold=True, size=9, color='FFFFFF')
    shade_cell(t_carga.rows[0].cells[j], AZUL_DARK)

for i, (resp, n, nivel, obs) in enumerate(carga):
    cell_text(t_carga.rows[i+1].cells[0], resp,  size=10, bold=True)
    cell_text(t_carga.rows[i+1].cells[1], str(n), size=12, bold=True,
              color=PRIORIDAD_TEXTO.get(nivel, '000000'), align='CENTER')
    cell_text(t_carga.rows[i+1].cells[2], nivel, size=9,  bold=True,
              color=PRIORIDAD_TEXTO.get(nivel, '000000'))
    cell_text(t_carga.rows[i+1].cells[3], obs,   size=9)
    col_bg = PRIORIDAD_COLOR.get(nivel, 'FFFFFF')
    for j in range(4):
        shade_cell(t_carga.rows[i+1].cells[j], col_bg)

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 6. RECOMENDACIONES
# ═══════════════════════════════════════════════════
titulo(doc, '6. Recomendaciones')
doc.add_paragraph()

recomendaciones = [
    ('Escalar bloqueos formalmente',
     'Todo bloqueo externo (META, migraciones, validaciones) debe escalarse por correo '
     'para garantizar trazabilidad y visibilidad ante la gerencia.'),
    ('Programar jornada Data Center',
     'Definir una fecha de sábado con reconocimiento económico para liquidar la deuda '
     'técnica del data center de una sola vez.'),
    ('Proteger bloques de trabajo estratégico',
     'Separar en la agenda semanal bloques de 2-3 horas para iniciativas estratégicas, '
     'blindados de interrupciones operativas.'),
    ('Completar empalme de conocimiento',
     'El empalme Johan Espino – Johan Velandia debe completarse esta semana. Documentar en repositorio '
     'de soporte interno para garantizar continuidad ante ausencias.'),
    ('Documentar todas las acciones en ClickUp',
     'Cada tarea ejecutada debe tener al menos un comentario con evidencia (captura, '
     'enlace, informe) antes de marcarse como completada.'),
]

for titulo_r, texto_r in recomendaciones:
    bullet(doc, texto_r, bold_label=titulo_r)

doc.add_paragraph()

# ── Pie ──
p_pie = doc.add_paragraph()
p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_pie = p_pie.add_run(
    'Documento generado automáticamente — Blend 360 Engineering Colombia | '
    'Centro de Servicios P0321 | Corte: 14 de abril de 2026')
r_pie.font.name = 'Montserrat'
r_pie.font.size = Pt(8)
r_pie.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.save(SALIDA)
print(f'Resumen ejecutivo generado: {SALIDA}')
